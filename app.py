import os
import math
import json
import re
import time
from datetime import datetime
from functools import wraps
from flask import Flask, request, jsonify, render_template, Response, make_response, redirect
from dotenv import load_dotenv
from flask_sqlalchemy import SQLAlchemy
import bcrypt
import jwt
import google.generativeai as genai
import secrets
import gzip
import io
import logging
from logging.handlers import RotatingFileHandler
from collections import defaultdict

# Load existing environment variables
load_dotenv()

app = Flask(__name__)

# --- SYSTEM ERROR LOGGING CONFIGURATION ---
if not os.path.exists('logs'):
    os.makedirs('logs')

log_handler = RotatingFileHandler('logs/app_errors.log', maxBytes=1000000, backupCount=3)
log_handler.setLevel(logging.WARNING)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
log_handler.setFormatter(formatter)
app.logger.addHandler(log_handler)

# --- RATE LIMITING MIDDLEWARE ---
RATE_LIMIT_TRACKER = defaultdict(list)

def rate_limit(limit=60, period=60):
    """Sliding-window IP rate limiter decorator."""
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            ip = request.remote_addr
            now = time.time()
            hits = [t for t in RATE_LIMIT_TRACKER[ip] if now - t < period]
            RATE_LIMIT_TRACKER[ip] = hits
            if len(hits) >= limit:
                app.logger.warning(f"Rate limit exceeded by IP: {ip} for path: {request.path}")
                return jsonify({'success': False, 'message': 'Too many requests. Please try again later.'}), 429
            RATE_LIMIT_TRACKER[ip].append(now)
            return f(*args, **kwargs)
        return wrapper
    return decorator

# --- CSRF PROTECTION MIDDLEWARE ---
CSRF_SECRET_KEY = os.environ.get("CSRF_SECRET_KEY", "universe_csrf_super_token_9918")

def generate_csrf_token():
    return secrets.token_hex(32)

@app.before_request
def global_csrf_protect():
    """Verify CSRF tokens globally for mutating API requests."""
    if request.method in ['POST', 'PUT', 'DELETE']:
        path = request.path
        exclusions = [
            'login', 'register', 'verify-email', 'forgot-password', 
            'reset-password', 'refresh', 'logout',
            'predict-colleges', 'find-scholarships', 'recommend', 'sync'
        ]
        if path.startswith('/api/') and not any(p in path for p in exclusions):
            # Bypass CSRF if a valid Bearer token is provided
            auth_header = request.headers.get('Authorization')
            if auth_header and auth_header.startswith('Bearer '):
                token = auth_header.split(' ')[1]
                from utils import verify_jwt
                if verify_jwt(token):
                    return
            
            cookie_token = request.cookies.get('csrf_token')
            header_token = request.headers.get('X-CSRF-Token')
            if not cookie_token or not header_token or cookie_token != header_token:
                app.logger.warning(f"CSRF verification failed for IP: {request.remote_addr} on path: {path}")
                return jsonify({'success': False, 'message': 'CSRF verification failed.'}), 403


# --- GZIP COMPRESSION & SECURITY HEADERS ---
@app.after_request
def apply_security_and_compression(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    
    # Caching headers for static assets
    if request.path.startswith('/static/'):
        response.headers['Cache-Control'] = 'public, max-age=31536000, immutable'
    
    csp = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' 'unsafe-eval' https://unpkg.com https://cdn.jsdelivr.net https://cdn.tailwindcss.com; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com https://cdnjs.cloudflare.com https://cdn.jsdelivr.net; "
        "font-src 'self' https://fonts.gstatic.com https://cdnjs.cloudflare.com; "
        "img-src * data:; "
        "connect-src 'self';"
    )
    if 'Content-Security-Policy' not in response.headers:
        response.headers['Content-Security-Policy'] = csp
        
    accept_encoding = request.headers.get('Accept-Encoding', '')
    if (response.status_code == 200 and
        'gzip' in accept_encoding.lower() and
        'Content-Encoding' not in response.headers and
        response.mimetype in ['text/html', 'text/css', 'application/javascript', 'application/json']):
        try:
            response.direct_passthrough = False
            data = response.get_data()
            gzip_buffer = io.BytesIO()
            with gzip.GzipFile(mode='wb', fileobj=gzip_buffer) as gzip_file_w:
                gzip_file_w.write(data)
            
            response.set_data(gzip_buffer.getvalue())
            response.headers['Content-Encoding'] = 'gzip'
            response.headers['Content-Length'] = len(response.get_data())
            response.headers['Vary'] = 'Accept-Encoding'
        except Exception as e:
            app.logger.warning(f"GZIP Compression failed: {e}")
            
    return response

# --- CENTRALIZED ERROR HANDLING ---
@app.errorhandler(404)
def handle_404(e):
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'message': 'Endpoint not found.'}), 404
    return redirect('/')

@app.errorhandler(500)
@app.errorhandler(Exception)
def handle_exception(e):
    try:
        db.session.rollback()
    except Exception:
        pass
    app.logger.error(f"Centralized Server Exception: {str(e)}", exc_info=True)
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'message': 'An internal server error occurred. Please try again later.'}), 500
    return make_response("<h3>An internal error occurred. Our engineers have been notified.</h3>", 500)


# JWT Configuration Parameters
SECRET_KEY = os.environ.get("JWT_SECRET_KEY", "universe_smart_secret_token_1882")

# SQL Database Configuration (PostgreSQL/custom URL if configured, else fallback to SQLite)
db_uri = os.environ.get("DATABASE_URL")
if db_uri:
    if db_uri.startswith("postgres"):
        db_uri = db_uri.replace("postgres://", "postgresql://", 1)
else:
    db_uri = "sqlite:///universe_portal.db"

app.config['SQLALCHEMY_DATABASE_URI'] = db_uri
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'connect_args': {
        'timeout': 30
    } if db_uri.startswith('sqlite') else {}
}

from models import db, User, University, Department, Course, Faculty, Announcement, GalleryItem, PlacementRecord, Scholarship, FAQItem, ChatSession, ChatHistoryItem, Bookmark, NotificationItem, SystemSetting, DocumentChunk, FavoriteCourse, Application, RecentSearch, RecentlyViewedItem, UserPreference, CommunityThread, CommunityPost, MockInterviewSession, ResumeProfile, UserSession, Hostel, Event, AdmissionChecklist
db.init_app(app)

from controllers.crud_factory import create_crud_blueprint
app.register_blueprint(create_crud_blueprint('universities_v2', University, auth_read=False, auth_write=True), url_prefix='/api/v2/universities')
app.register_blueprint(create_crud_blueprint('departments_v2', Department, auth_read=False, auth_write=True), url_prefix='/api/v2/departments')
app.register_blueprint(create_crud_blueprint('courses_v2', Course, auth_read=False, auth_write=True), url_prefix='/api/v2/courses')
app.register_blueprint(create_crud_blueprint('faculty_v2', Faculty, auth_read=False, auth_write=True), url_prefix='/api/v2/faculty')
app.register_blueprint(create_crud_blueprint('announcements_v2', Announcement, auth_read=False, auth_write=True), url_prefix='/api/v2/announcements')
app.register_blueprint(create_crud_blueprint('gallery_v2', GalleryItem, auth_read=False, auth_write=True), url_prefix='/api/v2/gallery')
app.register_blueprint(create_crud_blueprint('placements_v2', PlacementRecord, auth_read=False, auth_write=True), url_prefix='/api/v2/placements')
app.register_blueprint(create_crud_blueprint('scholarships_v2', Scholarship, auth_read=False, auth_write=True), url_prefix='/api/v2/scholarships')
app.register_blueprint(create_crud_blueprint('faqs_v2', FAQItem, auth_read=False, auth_write=True), url_prefix='/api/v2/faqs')
app.register_blueprint(create_crud_blueprint('chat_sessions_v2', ChatSession, auth_read=True, auth_write=True), url_prefix='/api/v2/chat-sessions')
app.register_blueprint(create_crud_blueprint('chat_history_v2', ChatHistoryItem, auth_read=True, auth_write=True), url_prefix='/api/v2/chat-history')
app.register_blueprint(create_crud_blueprint('bookmarks_v2', Bookmark, auth_read=True, auth_write=True), url_prefix='/api/v2/bookmarks')
app.register_blueprint(create_crud_blueprint('notifications_v2', NotificationItem, auth_read=True, auth_write=True), url_prefix='/api/v2/notifications')
app.register_blueprint(create_crud_blueprint('settings_v2', SystemSetting, auth_read=True, auth_write=True), url_prefix='/api/v2/settings')
app.register_blueprint(create_crud_blueprint('users_v2', User, auth_read=True, auth_write=True), url_prefix='/api/v2/users')
app.register_blueprint(create_crud_blueprint('document_chunks_v2', DocumentChunk, auth_read=True, auth_write=True), url_prefix='/api/v2/document-chunks')

# Hashing utilities using bcrypt
def hash_password(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def check_password(password_hash, password):
    try:
        return bcrypt.checkpw(password.encode('utf-8'), password_hash.encode('utf-8'))
    except Exception:
        return False

from utils import generate_jwt, verify_jwt, token_required, admin_required

# Constants
INFO_FILE = 'university_info.json'
UNIVERSE_FILE = 'universe_data.json'
USERS_FILE = 'users.json'
LOGS_FILE = 'chat_logs.json'
ANNOUNCEMENTS_FILE = 'announcements.json'

# --- File Loader Helpers ---

def load_universe_data():
    """Dynamically reconstruct universe data from relational SQL tables."""
    data = {}
    try:
        univs = University.query.all()
        for u in univs:
            courses = Course.query.filter_by(university_id=u.id).all()
            faculty = Faculty.query.filter_by(university_id=u.id).all()
            announcements = Announcement.query.filter_by(university_id=u.id).all()
            gallery = GalleryItem.query.filter_by(university_id=u.id).all()
            placement = PlacementRecord.query.filter_by(university_id=u.id).first()
            faqs = FAQItem.query.filter_by(university_id=u.id).all()
            
            fees_courses = []
            for c in courses:
                fees_courses.append({
                    'course_name': c.name,
                    'duration': c.duration or '4 Years',
                    'tuition_fee_per_year': c.fee
                })
            syllabus = {}
            for c in courses:
                if c.syllabus_json:
                    try:
                        syllabus[c.name] = json.loads(c.syllabus_json)
                    except Exception:
                        pass
                        
            top_rec = []
            if placement and placement.top_recruiters_json:
                try:
                    top_rec = json.loads(placement.top_recruiters_json)
                except Exception:
                    pass
                    
            courses_count = Course.query.filter_by(university_id=u.id).count()
            faculty_count = Faculty.query.filter_by(university_id=u.id).count()
            departments_count = Department.query.filter_by(university_id=u.id).count()
            students_count = User.query.filter_by(preferred_university=u.id).count()
            if students_count == 0:
                students_count = 1200 + (len(u.university_name) * 73) % 4000
                
            highest_pkg = placement.highest_package if placement else "N/A"
            avg_pkg = placement.average_package if placement else "N/A"
            rate = placement.placement_rate if placement else "N/A"
            
            stats = {
                'total_students': {'value': students_count, 'growth': '+4.8%', 'sparkline': [60, 65, 70, 75, 78, 80, 85]},
                'courses': {'value': courses_count or 12, 'growth': '+2.5%', 'sparkline': [10, 11, 12, 12, 13, 14, 15]},
                'faculty': {'value': faculty_count or 8, 'growth': '+3.1%', 'sparkline': [5, 6, 7, 7, 8, 9, 10]},
                'departments': {'value': departments_count or 3, 'growth': '+0.5%', 'sparkline': [2, 3, 3, 3, 3, 3, 3]},
                'placement_rate': {'value': rate if '%' in str(rate) else f"{rate}%" if rate != 'N/A' else '80.0%', 'growth': '+1.2%', 'sparkline': [70, 72, 75, 78, 80, 82, 85]},
                'average_package': {'value': avg_pkg, 'growth': '+5.8%', 'sparkline': [40, 42, 45, 48, 50, 52, 55]},
                'highest_package': {'value': highest_pkg, 'growth': '+12.5%', 'sparkline': [20, 25, 28, 30, 32, 35, 37]},
                'research_papers': {'value': 120 + (len(u.university_name) * 11) % 500, 'growth': '+8.5%', 'sparkline': [80, 90, 100, 110, 115, 120, 125]}
            }
            
            data[u.id] = {
                'university_name': u.university_name,
                'description': u.description or 'Explore course syllabus, annual fee sheets, and campus life.',
                'contact': {
                    'email': u.email or '',
                    'phone': u.phone or '',
                    'address': u.address or '',
                    'office_hours': u.office_hours or '',
                    'website': u.website or ''
                },
                'fees': {
                    'courses': fees_courses,
                    'payment_methods': ['Online NetBanking', 'Demand Draft'],
                    'installment_plan': 'Payable in semester installments.',
                    'refund_policy': 'As per institutional refund guidelines.'
                },
                'exams': {
                    'schedule': 'Sessional exams are conducted monthly. End-Sem in December and May.',
                    'policies': 'Minimum 75% attendance is mandatory to appear for exams.'
                },
                'syllabus': syllabus,
                'placements': {
                    'highest_package': highest_pkg,
                    'average_package': avg_pkg,
                    'placement_rate': rate,
                    'top_recruiters': top_rec
                },
                'hostel': {
                    'details': 'Hostels with basic amenities, study tables, and Wi-Fi.',
                    'fees': 'INR 25,000 per year',
                    'facilities': ['RO Drinking Water', '24/7 Security']
                },
                'campus_facilities': ['Central Library', 'Advanced Computing Labs', 'Sports pavilion'],
                'announcements': [{'title': a.title, 'type': a.type, 'desc': a.desc} for a in announcements],
                'events': [],
                'news': [],
                'faculty': [{'name': f.name, 'designation': f.designation, 'department': f.department, 'email': f.email} for f in faculty],
                'faqs': [{'question': faq.question, 'answer': faq.answer} for faq in faqs],
                'gallery': [g.image_url for g in gallery],
                'ranking': u.ranking or 'N/A',
                'accreditation': u.accreditation or 'N/A',
                'logo': u.logo or 'fa-solid fa-graduation-cap',
                'stats': stats,
                'brochure_url': u.brochure_url or '#',
                'virtual_tour_url': u.virtual_tour_url or '#'
            }
    except Exception as e:
        print(f"Error reading database for load_universe_data: {e}")
    return data

def save_universe_data(data):
    """Save/update entire universe dataset into relational SQL tables."""
    try:
        for u_id, u_info in data.items():
            univ = University.query.get(u_id)
            if not univ:
                univ = University(
                    id=u_id, 
                    university_name=u_info.get('university_name', u_id.upper()),
                    description=u_info.get('description', 'Explore course syllabus, annual fee sheets, and campus life.')
                )
                db.session.add(univ)
                db.session.flush()
                
            contact = u_info.get('contact', {})
            univ.description = u_info.get('description', univ.description or 'Explore course syllabus, annual fee sheets, and campus life.')
            univ.email = contact.get('email')
            univ.phone = contact.get('phone')
            univ.address = contact.get('address')
            univ.office_hours = contact.get('office_hours')
            univ.website = contact.get('website')
            univ.ranking = u_info.get('ranking')
            univ.accreditation = u_info.get('accreditation')
            univ.logo = u_info.get('logo')
            univ.brochure_url = u_info.get('brochure_url')
            univ.virtual_tour_url = u_info.get('virtual_tour_url')
            
            # Clear relations and rebuild
            Course.query.filter_by(university_id=u_id).delete()
            Faculty.query.filter_by(university_id=u_id).delete()
            Announcement.query.filter_by(university_id=u_id).delete()
            GalleryItem.query.filter_by(university_id=u_id).delete()
            PlacementRecord.query.filter_by(university_id=u_id).delete()
            FAQItem.query.filter_by(university_id=u_id).delete()
            
            # Courses
            courses_data = u_info.get("fees", {}).get("courses", [])
            syllabus_dict = u_info.get("syllabus", {})
            if isinstance(courses_data, list):
                for c_info in courses_data:
                    c_name = c_info.get("course_name") or c_info.get("name")
                    c_fee = c_info.get("tuition_fee_per_year") or c_info.get("fee")
                    c_duration = c_info.get("duration") or "4 Years"
                    
                    c_syll = syllabus_dict.get(c_name, {}) if isinstance(syllabus_dict, dict) else {}
                    db.session.add(Course(
                        university_id=u_id,
                        name=c_name,
                        duration=c_duration,
                        fee=c_fee,
                        syllabus_json=json.dumps(c_syll)
                    ))
            elif isinstance(courses_data, dict):
                for c_name, c_fee in courses_data.items():
                    c_syll = syllabus_dict.get(c_name, {}) if isinstance(syllabus_dict, dict) else {}
                    duration = "4 Years" if "B.Tech" in c_name or "B.Sc" in c_name else "2 Years"
                    if "BCA" in c_name or "BBA" in c_name:
                        duration = "3 Years"
                    db.session.add(Course(
                        university_id=u_id,
                        name=c_name,
                        duration=duration,
                        fee=c_fee,
                        syllabus_json=json.dumps(c_syll)
                    ))
                
            # Faculty
            for f_info in u_info.get('faculty', []):
                db.session.add(Faculty(
                    university_id=u_id,
                    name=f_info.get('name'),
                    designation=f_info.get('designation'),
                    department=f_info.get('department'),
                    email=f_info.get('email')
                ))
                
            # Announcements
            for a_info in u_info.get('announcements', []):
                db.session.add(Announcement(
                    university_id=u_id,
                    title=a_info.get('title'),
                    type=a_info.get('type', 'General'),
                    desc=a_info.get('desc')
                ))
                
            # Gallery
            for img_url in u_info.get('gallery', []):
                db.session.add(GalleryItem(
                    university_id=u_id,
                    image_url=img_url
                ))
                
            # Placements
            placements = u_info.get('placements', {})
            if placements:
                db.session.add(PlacementRecord(
                    university_id=u_id,
                    highest_package=placements.get('highest_package'),
                    average_package=placements.get('average_package'),
                    placement_rate=placements.get('placement_rate'),
                    top_recruiters_json=json.dumps(placements.get('top_recruiters', []))
                ))
                
            # FAQs
            for faq_info in u_info.get('faqs', []):
                db.session.add(FAQItem(
                    university_id=u_id,
                    question=faq_info.get('question'),
                    answer=faq_info.get('answer')
                ))
        db.session.commit()
        return True
    except Exception as e:
        print(f"Error saving universe data to DB: {e}")
        db.session.rollback()
        return False

def load_university_info():
    """Fallback helper. Reads parul university entry as standard info."""
    data = load_universe_data()
    return data.get("parul", {})

def save_university_info(new_data):
    """Fallback helper. Saves parul university entry."""
    data = load_universe_data()
    data["parul"] = new_data
    return save_universe_data(data)

def load_users():
    """Load persistent users list from JSON."""
    if not os.path.exists(USERS_FILE):
        return []
    try:
        with open(USERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading users: {e}")
        return []

def save_users(data):
    """Save users list to JSON."""
    try:
        with open(USERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error saving users: {e}")
        return False

def load_chat_logs():
    """Load helpdesk log registry from JSON."""
    if not os.path.exists(LOGS_FILE):
        return []
    try:
        with open(LOGS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading logs: {e}")
        return []

def save_chat_logs(data):
    """Save helpdesk logs registry to JSON."""
    try:
        with open(LOGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error saving logs: {e}")
        return False

def log_activity(user_id, action):
    """Log audit actions into chat_logs.json."""
    logs = load_chat_logs()
    logs.append({
        'user_id': user_id,
        'action': action,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    })
    save_chat_logs(logs)

def load_announcements():
    """Load announcements list from JSON."""
    if not os.path.exists(ANNOUNCEMENTS_FILE):
        return []
    try:
        with open(ANNOUNCEMENTS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading announcements: {e}")
        return []

def save_announcements(data):
    """Save announcements list to JSON."""
    try:
        with open(ANNOUNCEMENTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error saving announcements: {e}")
        return False

def update_env_variable(key, value):
    """Helper to update keys in .env file to persist configurations."""
    try:
        lines = []
        if os.path.exists('.env'):
            with open('.env', 'r', encoding='utf-8') as f:
                lines = f.readlines()
        
        updated = False
        new_lines = []
        for line in lines:
            if line.strip().startswith(f"{key}="):
                new_lines.append(f"{key}={value}\n")
                updated = True
            else:
                new_lines.append(line)
        
        if not updated:
            new_lines.append(f"{key}={value}\n")
            
        with open('.env', 'w', encoding='utf-8') as f:
            f.writelines(new_lines)
            
        # Update current process environment variable
        os.environ[key] = value
        return True
    except Exception as e:
        print(f"Error updating env file: {e}")
        return False

def get_api_key():
    """Get the Gemini API key from environment."""
    return os.environ.get("GEMINI_API_KEY", "").strip()

# --- Local NLP Matcher Engine ---
def match_course(query, courses_list):
    """Find if a specific course name is mentioned in the query."""
    query_lower = query.lower()
    for course in courses_list:
        name = course['course_name'].lower()
        keywords = [w for w in re.split(r'\W+', name) if len(w) > 2 and w not in ['and', 'bsc', 'tech', 'bba', 'mtech', 'mba', 'bachelor', 'master']]
        if name in query_lower or any(kw in query_lower for kw in keywords):
            return course
    return None

def local_nlp_chatbot(query, data):
    """
    Offline/local chatbot logic that parses the user query and matches
    against parameters of the selected university.
    """
    query_lower = query.lower()
    univ_name = data.get("university_name", "the University")
    
    # 1. Hello / Greeting checks
    if any(greet in query_lower for greet in ["hello", "hi", "hey", "greetings", "good morning", "good afternoon"]):
        return (
            f"Hello! Welcome to the **{univ_name}** Helpdesk. "
            "I can assist you with questions about:\n"
            "- 📥 **Admissions** (eligibility, deadlines, requirements)\n"
            "- 💰 **Fees & Installments** (tuition prices, refunds)\n"
            "- 📅 **Exam Schedules & Policies** (mid-terms, grading, attendance)\n"
            "- 📚 **Syllabus** (Engineering, Business, Science, etc.)\n\n"
            "What would you like to know today?"
        )
    
    # 2. General Contact Details
    if any(k in query_lower for k in ["contact", "phone", "email", "address", "call", "location", "office", "website"]):
        c = data.get("contact", {})
        return (
            f"Here are the contact details for **{univ_name}**:\n\n"
            f"- 📧 **Email:** {c.get('email', 'N/A')}\n"
            f"- 📞 **Phone:** {c.get('phone', 'N/A')}\n"
            f"- 📍 **Address:** {c.get('address', 'N/A')}\n"
            f"- 🕒 **Office Hours:** {c.get('office_hours', 'N/A')}\n"
            f"- 🌐 **Official Website:** [{c.get('website', 'Visit Link')}]({c.get('website', '#')})"
        )
        
    # 3. Admissions Queries
    if any(k in query_lower for k in ["admit", "admission", "apply", "deadline", "requirement", "eligible", "scholarship"]):
        ad = data.get("admissions", {})
        
        # Scholarship specific
        if "scholar" in query_lower:
            sc = ad.get("scholarships", {})
            return (
                f"### 🎓 Scholarships at {univ_name}\n\n"
                f"- **Merit-Based:** {sc.get('merit_based', 'N/A')}\n"
                f"- **Need-Based:** {sc.get('need_based', 'N/A')}\n"
                f"- **Application Deadline:** {sc.get('deadline', 'N/A')}\n\n"
                "You can apply for scholarships directly during the online admission process."
            )
            
        # Postgrad specific
        if any(p in query_lower for p in ["postgrad", "pg", "master", "m.tech", "mba", "phd"]):
            pg = ad.get("postgraduate", {})
            reqs = "\n".join([f"- {r}" for r in pg.get("requirements", [])])
            return (
                f"### 📥 Postgraduate Admissions\n\n"
                f"- **Eligibility:** {pg.get('eligibility', 'N/A')}\n"
                f"- **Deadline:** **{pg.get('deadline', 'N/A')}**\n"
                f"- **Requirements:**\n{reqs}\n"
                f"- **Admission Process:** {pg.get('process', 'N/A')}"
            )
            
        # Undergrad (default)
        ug = ad.get("undergraduate", {})
        reqs = "\n".join([f"- {r}" for r in ug.get("requirements", [])])
        return (
            f"### 📥 Undergraduate Admissions\n\n"
            f"- **Eligibility:** {ug.get('eligibility', 'N/A')}\n"
            f"- **Deadline:** **{ug.get('deadline', 'N/A')}**\n"
            f"- **Requirements:**\n{reqs}\n"
            f"- **Admission Process:** {ug.get('process', 'N/A')}\n\n"
            "*For Postgraduate admissions, please specify 'postgrad' or 'masters' in your query.*"
        )

    # 4. Fees Queries
    if any(k in query_lower for k in ["fee", "cost", "price", "tuition", "payment", "installment", "refund"]):
        fees = data.get("fees", {})
        courses = fees.get("courses", [])
        
        # Check if course-specific fee is requested
        matched_c = match_course(query_lower, courses)
        if matched_c:
            return (
                f"### 💰 Fee Structure for {matched_c['course_name']}\n\n"
                f"- **Level:** {matched_c['level']}\n"
                f"- **Duration:** {matched_c['duration']}\n"
                f"- **Tuition Fee:** **{matched_c['tuition_fee_per_year']}** per year\n\n"
                f"**Payment policy:** {fees.get('payment_methods', '')}"
            )
            
        # Refund policy specific
        if "refund" in query_lower:
            return (
                f"### 🔄 Refund Policy\n\n"
                f"{fees.get('refund_policy', 'N/A')}"
            )
            
        # Installment plan specific
        if any(ins in query_lower for ins in ["installment", "plan", "split", "parts"]):
            return (
                f"### 💳 Fee Installment Plan\n\n"
                f"{fees.get('installment_plan', 'N/A')}"
            )
            
        # General list of fees
        fee_table = "| Course Name | Level | Duration | Tuition (Per Year) |\n|---|---|---|---|\n"
        for c in courses:
            fee_table += f"| {c['course_name']} | {c['level']} | {c['duration']} | **{c['tuition_fee_per_year']}** |\n"
            
        return (
            f"### 💰 Tuition Fees Overview\n\n"
            f"{fee_table}\n"
            f"- **Payment Methods:** {fees.get('payment_methods', 'N/A')}\n"
            f"- **Installment Details:** {fees.get('installment_plan', 'N/A')}"
        )

    # 5. Exam Schedules & Policies
    if any(k in query_lower for k in ["exam", "schedule", "test", "midterm", "mid-term", "final", "grading", "attendance", "makeup", "make-up"]):
        ex = data.get("exams", {})
        sch = ex.get("schedule", {})
        pol = ex.get("policies", {})
        
        if any(p in query_lower for p in ["policy", "attendance", "grade", "makeup", "make-up"]):
            return (
                f"### 📅 Academic & Exam Policies\n\n"
                f"- 🏫 **Attendance Requirement:** {pol.get('attendance_requirement', 'N/A')}\n"
                f"- 📊 **Grading System:** {pol.get('grading_system', 'N/A')}\n"
                f"- 🤒 **Makeup Exams:** {pol.get('makeup_exams', 'N/A')}"
            )
            
        return (
            f"### 📅 Exam Schedules ({univ_name})\n\n"
            f"- 📝 **Mid-Term Exams:** {sch.get('mid_term', 'N/A')}\n"
            f"- 🎓 **Final Exams:** {sch.get('final_exam', 'N/A')}\n"
            f"- 📢 **Results Release Date:** {sch.get('results_release', 'N/A')}\n\n"
            "Make sure you maintain the mandatory attendance to sit for your exams!"
        )

    # 6. Syllabus Queries
    if any(k in query_lower for k in ["syllabus", "course", "subject", "curriculum", "semester", "learn", "study"]):
        sy = data.get("syllabus", {})
        
        matched_c_name = None
        for c_name in sy.keys():
            keywords = [w for w in re.split(r'\W+', c_name.lower()) if len(w) > 2 and w not in ['and', 'bsc', 'tech', 'bba', 'mtech', 'mba', 'bachelor', 'master']]
            if c_name.lower() in query_lower or any(kw in query_lower for kw in keywords):
                matched_c_name = c_name
                break
                
        if matched_c_name:
            course_syllabus = sy[matched_c_name]
            resp = f"### 📚 Syllabus: {matched_c_name}\n\n"
            for sem, subjects in course_syllabus.items():
                resp += f"**{sem}**:\n"
                for sub in subjects:
                    resp += f"- {sub}\n"
                resp += "\n"
            return resp
            
        course_names = "\n".join([f"- **{name}**" for name in sy.keys()])
        return (
            f"### 📚 Available Syllabi\n\n"
            "We have syllabus details for the following programs:\n"
            f"{course_names}\n\n"
            "Please ask: *'What is the syllabus for [Course Name]?'* to see the semester-wise subject breakdown."
        )

    # 7. Fallback
    return (
        f"I'm sorry, I couldn't find details matching your query about \"{query}\". "
        f"I can tell you about admissions, fees, exam schedule, syllabus, or contact details for **{univ_name}**.\n\n"
        "💡 *Tip: If you're an Admin, configure a **Gemini API Key** in the settings dashboard to unlock AI capabilities!*"
    )

# --- Pages Routes ---

@app.route('/')
def home():
    """Render landing page."""
    token = request.cookies.get('remember_token')
    if token:
        user_info = verify_jwt(token)
        if user_info:
            if user_info.get('role') == 'Admin':
                return redirect('/admin')
            return redirect('/portal')
    return render_template('index.html')

@app.route('/dashboard')
@app.route('/portal')
def portal():
    """Render Student Dashboard Portal."""
    token = request.args.get('token') or request.cookies.get('remember_token')
    if not token:
        # Server-side refresh token check for GET page loads
        ref_token = request.cookies.get('refresh_token')
        if ref_token:
            import time
            session_rec = UserSession.query.filter_by(refresh_token=ref_token).first()
            if session_rec and session_rec.expires_at >= int(time.time()):
                user = User.query.filter_by(email=session_rec.user_id).first()
                if user:
                    payload = {
                        'user_id': user.email,
                        'fullname': user.fullname,
                        'role': user.role,
                        'preferred_university': user.preferred_university or 'parul'
                    }
                    new_access_token = generate_jwt(payload, expires_in=900)
                    resp = make_response(render_template('portal.html'))
                    resp.set_cookie('remember_token', new_access_token, max_age=900, httponly=False, secure=not app.debug, samesite='Lax')
                    if not request.cookies.get('csrf_token'):
                        resp.set_cookie('csrf_token', generate_csrf_token(), max_age=86400 * 7, secure=not app.debug, samesite='Lax')
                    return resp
        return redirect('/')
    user_info = verify_jwt(token)
    if not user_info:
        resp = redirect('/')
        resp.delete_cookie('remember_token')
        resp.delete_cookie('refresh_token')
        return resp
    resp = make_response(render_template('portal.html'))
    resp.set_cookie('remember_token', token, max_age=900, httponly=False, secure=not app.debug, samesite='Lax')
    if not request.cookies.get('csrf_token'):
        resp.set_cookie('csrf_token', generate_csrf_token(), max_age=86400 * 7, secure=not app.debug, samesite='Lax')
    return resp

@app.route('/admin')
def admin_dashboard():
    """Render admin dashboard page."""
    token = request.args.get('token') or request.cookies.get('remember_token')
    if not token:
        ref_token = request.cookies.get('refresh_token')
        if ref_token:
            import time
            session_rec = UserSession.query.filter_by(refresh_token=ref_token).first()
            if session_rec and session_rec.expires_at >= int(time.time()):
                user = User.query.filter_by(email=session_rec.user_id).first()
                if user and user.role == 'Admin':
                    payload = {
                        'user_id': user.email,
                        'fullname': user.fullname,
                        'role': user.role,
                        'preferred_university': user.preferred_university or 'parul'
                    }
                    new_access_token = generate_jwt(payload, expires_in=900)
                    resp = make_response(render_template('admin.html'))
                    resp.set_cookie('remember_token', new_access_token, max_age=900, httponly=False, secure=not app.debug, samesite='Lax')
                    if not request.cookies.get('csrf_token'):
                        resp.set_cookie('csrf_token', generate_csrf_token(), max_age=86400 * 7, secure=not app.debug, samesite='Lax')
                    return resp
        return redirect('/')
    user_info = verify_jwt(token)
    if not user_info or user_info.get('role') != 'Admin':
        resp = redirect('/')
        resp.delete_cookie('remember_token')
        resp.delete_cookie('refresh_token')
        return resp
    resp = make_response(render_template('admin.html'))
    resp.set_cookie('remember_token', token, max_age=900, httponly=False, secure=not app.debug, samesite='Lax')
    if not request.cookies.get('csrf_token'):
        resp.set_cookie('csrf_token', generate_csrf_token(), max_age=86400 * 7, secure=not app.debug, samesite='Lax')
    return resp

# --- API Interaction Endpoints ---

@app.route('/api/chat-sessions', methods=['GET'])
@token_required
def get_chat_sessions():
    user_id = request.user['user_id']
    q = request.args.get('q', '').strip()
    try:
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 15))
    except ValueError:
        page = 1
        limit = 15
    offset = (page - 1) * limit
    
    query = ChatSession.query.filter_by(user_id=user_id)
    if q:
        from sqlalchemy import or_
        matching_sess_ids = [item.session_id for item in ChatHistoryItem.query.filter(ChatHistoryItem.user_id == user_id, ChatHistoryItem.message.ilike(f'%{q}%')).all()]
        query = query.filter(or_(ChatSession.title.ilike(f'%{q}%'), ChatSession.id.in_(matching_sess_ids)))
        
    total = query.count()
    sessions = query.order_by(ChatSession.id.desc()).offset(offset).limit(limit).all()
    
    resp = jsonify([s.to_dict() for s in sessions])
    resp.headers['X-Total-Count'] = str(total)
    resp.headers['X-Page'] = str(page)
    resp.headers['X-Limit'] = str(limit)
    return resp

@app.route('/api/chat-sessions', methods=['POST'])
@token_required
def create_chat_session():
    user_id = request.user['user_id']
    body = request.get_json() or {}
    univ_id = body.get('university_id', 'parul').strip()
    title = body.get('title', 'New Chat').strip()
    
    session = ChatSession(
        user_id=user_id,
        university_id=univ_id,
        title=title,
        timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    )
    db.session.add(session)
    db.session.commit()
    return jsonify({'success': True, 'session': session.to_dict()}), 201

@app.route('/api/chat-sessions/<int:session_id>', methods=['PUT'])
@token_required
def rename_chat_session(session_id):
    user_id = request.user['user_id']
    session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
    if not session:
        return jsonify({'success': False, 'message': 'Chat session not found'}), 404
        
    body = request.get_json() or {}
    title = body.get('title', '').strip()
    if not title:
        return jsonify({'success': False, 'message': 'Title is required'}), 400
        
    session.title = title
    db.session.commit()
    return jsonify({'success': True, 'session': session.to_dict()})

@app.route('/api/chat-sessions/<int:session_id>', methods=['DELETE'])
@token_required
def delete_chat_session(session_id):
    user_id = request.user['user_id']
    session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
    if not session:
        return jsonify({'success': False, 'message': 'Chat session not found'}), 404
        
    db.session.delete(session)
    db.session.commit()
    return jsonify({'success': True, 'message': 'Chat session deleted successfully'})

@app.route('/api/chat-sessions/<int:session_id>/messages', methods=['GET'])
@token_required
def get_session_messages(session_id):
    user_id = request.user['user_id']
    session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
    if not session:
        return jsonify({'success': False, 'message': 'Chat session not found'}), 404
        
    try:
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 50))
    except ValueError:
        page = 1
        limit = 50
    offset = (page - 1) * limit
    total = ChatHistoryItem.query.filter_by(session_id=session_id).count()
    messages = ChatHistoryItem.query.filter_by(session_id=session_id).order_by(ChatHistoryItem.id.asc()).offset(offset).limit(limit).all()
    
    resp = jsonify([m.to_dict() for m in messages])
    resp.headers['X-Total-Count'] = str(total)
    resp.headers['X-Page'] = str(page)
    resp.headers['X-Limit'] = str(limit)
    return resp

_GEMINI_API_BLOCKED = False

def is_valid_api_key(api_key):
    if not api_key:
        return False
    ak_lower = api_key.lower()
    if '•' in api_key or '*' in api_key or 'your_' in ak_lower or 'key_here' in ak_lower or 'dummy' in ak_lower or len(api_key) < 15:
        return False
    return True

def get_embedding(text, api_key=None):
    global _GEMINI_API_BLOCKED
    if _GEMINI_API_BLOCKED or not is_valid_api_key(api_key):
        return None
    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        result = genai.embed_content(
            model="models/embedding-001",
            content=text,
            task_type="retrieval_document",
            request_options={'timeout': 2.0}
        )
        return result['embedding']
    except Exception as e:
        print(f"Error generating embedding via Gemini API: {e}")
        if "plugin_credentials" in str(e) or "Illegal header" in str(e) or "metadata" in str(e) or "503" in str(e) or "timeout" in str(e).lower() or "deadline" in str(e).lower():
            _GEMINI_API_BLOCKED = True
            print("Gemini API is blocked/misconfigured in this environment. Bypassing subsequent calls.")
        return None

def cosine_similarity(v1, v2):
    dot = sum(a*b for a, b in zip(v1, v2))
    n1 = math.sqrt(sum(a*a for a in v1))
    n2 = math.sqrt(sum(b*b for b in v2))
    return dot / (n1 * n2) if n1 and n2 else 0.0

def retrieve_relevant_chunks(university_id, query_text, limit=3):
    """Retrieve top matched document chunks for a university using RAG embeddings."""
    global _GEMINI_API_BLOCKED
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    query_vector = None
    if is_valid_api_key(api_key) and not _GEMINI_API_BLOCKED:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            result = genai.embed_content(
                model="models/embedding-001",
                content=query_text,
                task_type="retrieval_query",
                request_options={'timeout': 2.0}
            )
            query_vector = result['embedding']
        except Exception as e:
            print(f"Error embedding query: {e}")
            if "plugin_credentials" in str(e) or "Illegal header" in str(e) or "metadata" in str(e) or "503" in str(e) or "timeout" in str(e).lower() or "deadline" in str(e).lower():
                _GEMINI_API_BLOCKED = True
                print("Gemini API is blocked/misconfigured in this environment. Bypassing subsequent query embeddings.")
            
    chunks = DocumentChunk.query.filter_by(university_id=university_id).all()
    
    scored_chunks = []
    if query_vector:
        for chunk in chunks:
            if chunk.embedding_json:
                chunk_vector = json.loads(chunk.embedding_json)
                score = cosine_similarity(query_vector, chunk_vector)
                scored_chunks.append((score, chunk))
    
    if not scored_chunks:
        def stem(w):
            if len(w) <= 3: return w
            if w.endswith('ies'): return w[:-3] + 'y'
            if w.endswith('es') and not w.endswith('aes') and not w.endswith('ees') and not w.endswith('oes'): return w[:-2]
            if w.endswith('s') and not w.endswith('ss'): return w[:-1]
            if w.endswith('ing'): return w[:-3]
            if w.endswith('ed'): return w[:-2]
            return w

        query_words = {stem(w) for w in re.findall(r'\w+', query_text.lower())}
        for chunk in chunks:
            chunk_words = {stem(w) for w in re.findall(r'\w+', chunk.content.lower())}
            intersection = query_words.intersection(chunk_words)
            score = len(intersection) / len(query_words) if query_words else 0.0
            scored_chunks.append((score, chunk))
            
    scored_chunks.sort(key=lambda x: x[0], reverse=True)
    return [chunk for score, chunk in scored_chunks[:limit]]

@app.route('/api/chat/stream', methods=['GET'])
def chat_stream():
    """Stream chatbot responses in real-time using Server-Sent Events (SSE)."""
    token = request.args.get('token')
    if not token:
        return "data: " + json.dumps({'error': 'Unauthorized: Token is missing!'}) + "\n\n", 401
        
    payload = verify_jwt(token)
    if not payload:
        return "data: " + json.dumps({'error': 'Unauthorized: Token is invalid!'}) + "\n\n", 401
        
    user_email = payload['user_id']
    
    session_id = request.args.get('session_id', type=int)
    message = request.args.get('message', '').strip()
    university_id = request.args.get('university_id', 'parul').strip()
    mode = request.args.get('mode', '').strip()
    
    if not session_id or not message:
        return "data: " + json.dumps({'error': 'Missing session_id or message parameter'}) + "\n\n", 400
        
    session = ChatSession.query.filter_by(id=session_id, user_id=user_email).first()
    if not session:
        return "data: " + json.dumps({'error': 'Chat session not found'}) + "\n\n", 404

    univ = University.query.get(university_id)
    univ_name = univ.university_name if univ else university_id.upper()
    
    univ_data = {}
    try:
        courses = Course.query.filter_by(university_id=university_id).all()
        fees_courses = [{"course_name": c.name, "tuition_fee_per_year": c.fee} for c in courses]
        faculty = Faculty.query.filter_by(university_id=university_id).all()
        announcements = Announcement.query.filter_by(university_id=university_id).all()
        placement = PlacementRecord.query.filter_by(university_id=university_id).first()
        faqs = FAQItem.query.filter_by(university_id=university_id).all()
        
        univ_data = {
            'university_name': univ_name,
            'fees': {'courses': fees_courses},
            'placements': {
                'highest_package': placement.highest_package if placement else 'N/A',
                'average_package': placement.average_package if placement else 'N/A',
                'placement_rate': placement.placement_rate if placement else 'N/A'
            },
            'announcements': [{'title': a.title, 'desc': a.desc} for a in announcements],
            'faculty': [{'name': f.name, 'designation': f.designation, 'department': f.department} for f in faculty],
            'faqs': [{'question': faq.question, 'answer': faq.answer} for faq in faqs]
        }
    except Exception as e:
        print(f"Error loading university details context: {e}")

    # Fetch relevant document chunks using RAG vector database matching
    rag_chunks = []
    try:
        rag_chunks = retrieve_relevant_chunks(university_id, message, limit=3)
    except Exception as ex:
        print(f"RAG Retrieval error: {ex}")
        
    rag_context_str = "\n\n".join([f"[Category: {c.category.upper()}]\n{c.content}" for c in rag_chunks])

    if mode == 'copilot':
        system_prompt = (
            f"You are the official UniVerse AI Admissions Copilot for {univ_name}.\n\n"
            f"Retrieved University Knowledge Base Context:\n"
            f"{rag_context_str}\n\n"
            f"University Context Profile:\n{json.dumps(univ_data, indent=2)}\n\n"
            f"Admissions Guidelines:\n"
            f"1. Explain eligibility criteria, document requirements, and admission progress checkpoints clearly.\n"
            f"2. Suggest immediate next steps for application and admission schedules.\n"
            f"3. Recommend eligible scholarships, and list dynamic checkpoints / reminders with [ ] or [x] checkboxes.\n"
            f"4. Recommend next admission steps step-by-step.\n"
            f"5. Focus only on {univ_name} info. If asked about others, guide context switching.\n"
            f"6. Professional Markdown layout with checklist check-boxes [ ] or [x], bold headers, tables, etc."
        )
    else:
        system_prompt = (
            f"You are the official UniVerse AI Admission Assistant for {univ_name}.\n\n"
            f"Retrieved University Knowledge Base Context:\n"
            f"{rag_context_str}\n\n"
            f"University Context Profile:\n{json.dumps(univ_data, indent=2)}\n\n"
            f"Guidelines:\n"
            f"1. Prioritize answering using the Retrieved University Knowledge Base Context. If it contains details addressing the query, quote or summarize them accurately.\n"
            f"2. Help the student with admissions, fee structures, syllabus, placements, faculty, announcements, hostels, mess, research, and campus facilities.\n"
            f"3. Only discuss details about {univ_name}. If asked about other colleges, guide them to switch context.\n"
            f"4. Support standard conversational queries (greetings, 'thank you', 'who are you', etc.).\n"
            f"5. Format the output professionally using standard Markdown list bullet points, bolding, and markdown tables where appropriate.\n"
            f"6. Keep the tone friendly, academic, and direct."
        )

    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else None
    if not api_key:
        api_key = os.environ.get("GEMINI_API_KEY")

    def generate_events():
        try:
            with app.app_context():
                user_msg = ChatHistoryItem(
                    session_id=session_id,
                    user_id=user_email,
                    university_id=university_id,
                    sender='user',
                    message=message,
                    timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                )
                db.session.add(user_msg)
                
                active_session = ChatSession.query.get(session_id)
                if active_session and active_session.title == "New Chat":
                    active_session.title = message[:30] + ("..." if len(message) > 30 else "")
                db.session.commit()
        except Exception as ex:
            print(f"Error saving user message in generator: {ex}")

        bot_response_text = ""
        if api_key:
            try:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(
                    model_name='gemini-1.5-flash',
                    system_instruction=system_prompt
                )
                response = model.generate_content(message, stream=True)
                for chunk in response:
                    text_chunk = chunk.text
                    bot_response_text += text_chunk
                    yield f"data: {json.dumps({'text': text_chunk})}\n\n"
            except Exception as e:
                print(f"Gemini Streaming Error: {e}, falling back to local NLP")
                fallback_text = local_nlp_chatbot(message, univ_data)
                words = fallback_text.split(" ")
                for word in words:
                    time.sleep(0.05)
                    text_chunk = word + " "
                    bot_response_text += text_chunk
                    yield f"data: {json.dumps({'text': text_chunk})}\n\n"
        else:
            fallback_text = local_nlp_chatbot(message, univ_data)
            words = fallback_text.split(" ")
            for word in words:
                time.sleep(0.05)
                text_chunk = word + " "
                bot_response_text += text_chunk
                yield f"data: {json.dumps({'text': text_chunk})}\n\n"
                
        try:
            with app.app_context():
                bot_msg = ChatHistoryItem(
                    session_id=session_id,
                    user_id=user_email,
                    university_id=university_id,
                    sender='bot',
                    message=bot_response_text,
                    timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                )
                db.session.add(bot_msg)
                db.session.commit()
        except Exception as ex:
            print(f"Error saving bot message in generator: {ex}")

    return Response(generate_events(), mimetype='text/event-stream')

@app.route('/api/chat', methods=['POST'])
@rate_limit(limit=30, period=60)
@token_required
def chat():
    """Handle chat messages, routing to Gemini or Local NLP, and logs results."""
    start_time = time.time()
    body = request.get_json() or {}
    message = body.get('message', '').strip()
    university_id = body.get('university_id', 'parul').strip()
    session_id = body.get('session_id')
    
    if not message:
        return jsonify({'response': 'Please enter a message.'}), 400
        
    user_email = request.user['user_id']
    if session_id:
        try:
            user_msg = ChatHistoryItem(
                session_id=session_id,
                user_id=user_email,
                university_id=university_id,
                sender='user',
                message=message,
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )
            db.session.add(user_msg)
            
            active_session = ChatSession.query.get(session_id)
            if active_session and active_session.title == "New Chat":
                active_session.title = message[:30] + ("..." if len(message) > 30 else "")
            db.session.commit()
        except Exception as ex:
            print(f"Error saving user message: {ex}")
        
    universe_data = load_universe_data()
    univ_data = universe_data.get(university_id, universe_data.get("parul", {}))
    api_key = get_api_key()
    
    response_text = ""
    engine_source = ""
    status = "Success"
    error_msg = ""
    
    if api_key:
        try:
            # Configure Gemini API
            genai.configure(api_key=api_key)
            
            system_prompt = (
                f"You are the AI Student Helpdesk Chatbot for {univ_data.get('university_name', 'our university')}.\n"
                f"Your goal is to answer queries regarding admissions, fees, exams, syllabus, and contact info.\n\n"
                f"Official University Data Context:\n"
                f"{json.dumps(univ_data, indent=2)}\n\n"
                f"Guidelines:\n"
                f"1. Answer accurate facts using ONLY the database context provided above.\n"
                f"2. If the info is not in the database, politely state that you don't have this detail, "
                f"and advise them to contact the helpdesk at {univ_data.get('contact', {}).get('email', 'the admissions office')}.\n"
                f"3. Support standard conversational queries (greetings, 'thank you', 'who are you', etc.).\n"
                f"4. Format the output professionally using standard Markdown list bullet points, bolding, and markdown tables where appropriate.\n"
                f"5. Keep the tone friendly, academic, and direct."
            )
            
            model = genai.GenerativeModel(
                model_name='gemini-1.5-flash',
                system_instruction=system_prompt
            )
            
            response = model.generate_content(message)
            response_text = response.text
            engine_source = f"Gemini AI ({univ_data.get('university_name')})"
            
        except Exception as e:
            # Fall back to local NLP on API error
            print(f"Gemini API Error, falling back to Local NLP: {e}")
            response_text = local_nlp_chatbot(message, univ_data)
            engine_source = "Local NLP Engine (API Failover)"
            status = "Warning"
            error_msg = str(e)
    else:
        response_text = local_nlp_chatbot(message, univ_data)
        engine_source = f"Local NLP ({univ_data.get('university_name')})"

    latency = f"{time.time() - start_time:.2f}s"
    
    log_entry = {
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'user': request.user['user_id'],
        'query': message[:100] + ('...' if len(message) > 100 else ''),
        'engine': engine_source,
        'status': status,
        'latency': latency
    }
    logs = load_chat_logs()
    logs.insert(0, log_entry)
    save_chat_logs(logs[:200])

    if session_id:
        try:
            bot_msg = ChatHistoryItem(
                session_id=session_id,
                user_id=user_email,
                university_id=university_id,
                sender='bot',
                message=response_text,
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )
            db.session.add(bot_msg)
            db.session.commit()
        except Exception as ex:
            print(f"Error saving bot message: {ex}")

    return jsonify({
        'response': response_text,
        'source': engine_source,
        'latency': latency
    })

LAST_SYNCED = {}

def run_background_sync(app_obj, univ_id):
    """Worker function to run RAG crawler sync in a background thread with application context."""
    with app_obj.app_context():
        try:
            from services.realtime_fetcher import scrape_university_details
            scrape_university_details(univ_id)
        except Exception as e:
            db.session.rollback()
            print(f"Error in background sync for {univ_id}: {e}")

@app.route('/api/universities/<university_id>/sync', methods=['POST'])
def sync_university_data(university_id):
    """Trigger real-time crawling and database update for the selected university."""
    from services.realtime_fetcher import scrape_university_details
    try:
        success, message = scrape_university_details(university_id)
        if success:
            LAST_SYNCED[university_id] = time.time()
            return jsonify({'success': True, 'message': message})
        return jsonify({'success': False, 'message': message}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f"Sync failed: {str(e)}"}), 500


@app.route('/api/universities', methods=['GET', 'POST', 'DELETE'])
def universities():
    """Retrieve lists of universities, single university details, or add new universities."""
    if request.method != 'GET':
        token = None
        auth_header = request.headers.get('Authorization')
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]
        if not token:
            token = request.cookies.get('remember_token')
            
        if not token:
            return jsonify({'success': False, 'message': 'Authorization token is missing.'}), 401
            
        user_info = verify_jwt(token)
        if not user_info or user_info.get('role') != 'Admin':
            return jsonify({'success': False, 'message': 'Access denied.'}), 403

    if request.method == 'GET':
        univ_id = request.args.get('id')
        if univ_id:
            univ = University.query.get(univ_id)
            if not univ:
                return jsonify({'error': 'University not found'}), 404
            
            # Trigger background sync if not synced in last 1 hour
            now = time.time()
            if univ_id not in LAST_SYNCED or (now - LAST_SYNCED[univ_id]) > 3600:
                LAST_SYNCED[univ_id] = now
                import threading
                threading.Thread(target=run_background_sync, args=(app, univ_id)).start()
            
            courses = Course.query.filter_by(university_id=univ_id).all()
            faculty = Faculty.query.filter_by(university_id=univ_id).all()
            announcements = Announcement.query.filter_by(university_id=univ_id).all()
            gallery = GalleryItem.query.filter_by(university_id=univ_id).all()
            placement = PlacementRecord.query.filter_by(university_id=univ_id).first()
            faqs = FAQItem.query.filter_by(university_id=univ_id).all()
            
            fees_courses = []
            for c in courses:
                fees_courses.append({
                    'id': c.id,
                    'course_name': c.name,
                    'duration': c.duration or '4 Years',
                    'tuition_fee_per_year': c.fee
                })
            syllabus = {}
            for c in courses:
                if c.syllabus_json:
                    try:
                        syllabus[c.name] = json.loads(c.syllabus_json)
                    except Exception:
                        pass
                        
            top_rec = []
            if placement and placement.top_recruiters_json:
                try:
                    top_rec = json.loads(placement.top_recruiters_json)
                except Exception:
                    pass
                    
            courses_count = Course.query.filter_by(university_id=univ_id).count()
            faculty_count = Faculty.query.filter_by(university_id=univ_id).count()
            departments_count = Department.query.filter_by(university_id=univ_id).count()
            scholarships = Scholarship.query.filter_by(university_id=univ_id).all()
            # Retrieve realistic stats from universe_data.json if present
            real_students = 0
            real_papers = 0
            if os.path.exists(UNIVERSE_FILE):
                try:
                    with open(UNIVERSE_FILE, 'r', encoding='utf-8') as f:
                        u_data = json.load(f)
                        univ_stats = u_data.get(univ_id, {}).get('stats', {})
                        real_students = univ_stats.get('total_students', {}).get('value', 0)
                        real_papers = univ_stats.get('research_papers', {}).get('value', 0)
                except Exception:
                    pass
            
            if not real_students:
                real_students = 8500 + (len(univ.university_name) * 123) % 15000
            if not real_papers:
                real_papers = 450 + (len(univ.university_name) * 47) % 1200
                
            highest_pkg = placement.highest_package if placement else "N/A"
            avg_pkg = placement.average_package if placement else "N/A"
            rate = placement.placement_rate if placement else "N/A"
            
            extra = {}
            if univ.crawled_details_json:
                try:
                    extra = json.loads(univ.crawled_details_json)
                except Exception:
                    pass

            c_students = real_students or extra.get('total_students')
            c_faculty_count = extra.get('total_faculty') or faculty_count
            c_departments_count = extra.get('departments_count') or departments_count
            
            c_admission_dates = extra.get('admission_dates') or 'Information currently unavailable'
            c_application_deadlines = extra.get('application_deadlines') or 'Information currently unavailable'
            c_hostel_fees = extra.get('hostel_fees') or 'Information currently unavailable'
            c_hostel_facilities = extra.get('hostel_facilities') or []
            c_campus_facilities = extra.get('campus_facilities') or []
            c_research_centers = extra.get('research_centers') or []
            c_academic_calendar = extra.get('academic_calendar') or 'Information currently unavailable'
            c_source_attribution = extra.get('source_attribution') or 'Official University Website'
            
            stats = {
                'total_students': {'value': c_students or 'Information currently unavailable', 'growth': '+4.8%', 'sparkline': [60, 65, 70, 75, 78, 80, 85]},
                'courses': {'value': courses_count or 'Information currently unavailable', 'growth': '+2.5%', 'sparkline': [10, 11, 12, 12, 13, 14, 15]},
                'faculty': {'value': c_faculty_count or 'Information currently unavailable', 'growth': '+3.1%', 'sparkline': [5, 6, 7, 7, 8, 9, 10]},
                'departments': {'value': c_departments_count or 'Information currently unavailable', 'growth': '+0.5%', 'sparkline': [2, 3, 3, 3, 3, 3, 3]},
                'placement_rate': {'value': rate if '%' in str(rate) else f"{rate}%" if rate != 'N/A' else 'Information currently unavailable', 'growth': '+1.2%', 'sparkline': [70, 72, 75, 78, 80, 82, 85]},
                'average_package': {'value': avg_pkg, 'growth': '+5.8%', 'sparkline': [40, 42, 45, 48, 50, 52, 55]},
                'highest_package': {'value': highest_pkg, 'growth': '+12.5%', 'sparkline': [20, 25, 28, 30, 32, 35, 37]},
                'research_papers': {'value': real_papers or 'Information currently unavailable', 'growth': '+8.5%', 'sparkline': [80, 90, 100, 110, 115, 120, 125]}
            }
            
            hostel_db = Hostel.query.filter_by(university_id=univ_id).first()
            hostel_details = 'Information currently unavailable'
            hostel_fees_val = 'Information currently unavailable'
            hostel_facilities_list = ['Information currently unavailable']
            hostel_names = 'Separate Boys & Girls Hostels'
            hostel_warden = '9876543210'
            
            if hostel_db:
                hostel_dict = hostel_db.to_dict()
                hostel_details = hostel_db.details or 'AC and Non-AC premium hosteling options with 24/7 security, high-speed Wi-Fi, and modular laundry hubs.'
                hostel_fees_val = hostel_db.fees or 'N/A'
                hostel_facilities_list = hostel_dict.get('facilities') or ['24/7 High-speed Internet', 'Hygienic vegetarian mess', 'Common Gym', 'Power Backup']
                hostel_names = hostel_db.names or 'Separate Boys & Girls Hostels'
                hostel_warden = hostel_db.warden_contact or '9876543210'
            elif c_hostel_fees:
                hostel_details = 'AC and Non-AC hostel blocks with laundry services, study halls, and common dining halls.'
                hostel_fees_val = c_hostel_fees
                hostel_facilities_list = c_hostel_facilities if c_hostel_facilities else ['Wi-Fi', 'Common Mess']

            events_db = Event.query.filter_by(university_id=univ_id).all()
            events_list = []
            if events_db:
                events_list = [{'title': ev.title, 'date': ev.date or '', 'desc': ev.desc or ''} for ev in events_db]
            else:
                events_list = [
                    {"title": "Annual Cultural Fest", "date": "March 15, 2026", "desc": "Inter-university talent show and cultural gala."},
                    {"title": "TechExpo 2026", "date": "April 08, 2026", "desc": "Exhibition of top student projects and research papers."}
                ]

            resp_dict = {
                'id': univ.id,
                'university_name': univ.university_name,
                'description': univ.description or 'Explore course syllabus, annual fee sheets, and campus life.',
                'contact': {
                    'email': univ.email or 'Information currently unavailable',
                    'phone': univ.phone or 'Information currently unavailable',
                    'address': univ.address or 'Information currently unavailable',
                    'office_hours': univ.office_hours or 'Information currently unavailable',
                    'website': univ.website or ''
                },
                'admissions': {
                    'undergraduate': {
                        'eligibility': 'Pass in Class 12 exams with minimum marks and entrance score.',
                        'deadline': c_application_deadlines,
                        'requirements': [
                            'Online registration scorecard',
                            'Class 10 and 12 mark sheets',
                            'Transfer & Migration Certificate'
                        ],
                        'process': 'Register online, verify documents, lock choices, report to campus.'
                    },
                    'dates': c_admission_dates
                },
                'fees': {
                    'courses': fees_courses,
                    'payment_methods': ['Online NetBanking', 'Demand Draft'],
                    'installment_plan': 'Payable in semester installments.',
                    'refund_policy': 'As per institutional refund guidelines.'
                },
                'exams': {
                    'schedule': c_academic_calendar,
                    'policies': 'Minimum 75% attendance is mandatory to appear for exams.'
                },
                'syllabus': syllabus,
                'placements': {
                    'highest_package': highest_pkg,
                    'average_package': avg_pkg,
                    'placement_rate': rate,
                    'top_recruiters': top_rec
                },
                'hostel': {
                    'details': hostel_details,
                    'fees': hostel_fees_val,
                    'facilities': hostel_facilities_list,
                    'names': hostel_names,
                    'warden_contact': hostel_warden
                },
                'campus_facilities': c_campus_facilities if c_campus_facilities else ['Information currently unavailable'],
                'research_centers': c_research_centers if c_research_centers else ['Information currently unavailable'],
                'announcements': [{'id': a.id, 'title': a.title, 'type': a.type, 'desc': a.desc} for a in announcements],
                'events': events_list,
                'news': [],
                'faculty': [{'id': f.id, 'name': f.name, 'designation': f.designation, 'department': f.department, 'email': f.email} for f in faculty],
                'faqs': [{'question': faq.question, 'answer': faq.answer} for faq in faqs],
                'gallery': [{'id': g.id, 'image_url': g.image_url} for g in gallery],
                'ranking': univ.ranking or 'Information currently unavailable',
                'accreditation': univ.accreditation or 'Information currently unavailable',
                'logo': univ.logo or 'fa-solid fa-graduation-cap',
                'stats': stats,
                'scholarships': [{'id': s.id, 'title': s.title, 'eligibility': s.eligibility, 'amount': s.amount} for s in scholarships],
                'brochure_url': univ.brochure_url or '#',
                'virtual_tour_url': univ.virtual_tour_url or '#',
                'last_updated': univ.last_updated or '',
                'source_attribution': c_source_attribution
            }
            return jsonify(resp_dict)
            
        listings = []
        univs = University.query.all()
        for u in univs:
            listings.append({
                'id': u.id,
                'name': u.university_name,
                'location': u.address or 'N/A',
                'logo': u.logo or 'fa-solid fa-graduation-cap',
                'ranking': u.ranking or 'N/A',
                'accreditation': u.accreditation or 'N/A',
                'website': u.website or '#'
            })
        return jsonify(listings)
    elif request.method == 'POST':
        body = request.get_json() or {}
        univ_id = body.get('id', '').strip().lower()
        univ_details = body.get('details', {})
        
        if not univ_id or not univ_details:
            return jsonify({'success': False, 'message': 'University ID and details are required.'}), 400
            
        # Sync to universe_data.json
        data = load_universe_data()
        data[univ_id] = univ_details
        save_universe_data(data)
        
        univ = University.query.get(univ_id)
        if not univ:
            univ = University(
                id=univ_id, 
                university_name=univ_details.get('university_name', univ_id.upper()),
                description=univ_details.get('description', 'Explore course syllabus, annual fee sheets, and campus life.')
            )
            db.session.add(univ)
            db.session.flush()
            
        contact = univ_details.get('contact', {})
        univ.description = univ_details.get('description', univ.description or 'Explore course syllabus, annual fee sheets, and campus life.')
        univ.email = contact.get('email')
        univ.phone = contact.get('phone')
        univ.address = contact.get('address')
        univ.office_hours = contact.get('office_hours')
        univ.website = contact.get('website')
        univ.ranking = univ_details.get('ranking')
        univ.accreditation = univ_details.get('accreditation')
        univ.logo = univ_details.get('logo')
        univ.brochure_url = univ_details.get('brochure_url')
        univ.virtual_tour_url = univ_details.get('virtual_tour_url')
        
        Course.query.filter_by(university_id=univ_id).delete()
        Faculty.query.filter_by(university_id=univ_id).delete()
        Announcement.query.filter_by(university_id=univ_id).delete()
        GalleryItem.query.filter_by(university_id=univ_id).delete()
        PlacementRecord.query.filter_by(university_id=univ_id).delete()
        FAQItem.query.filter_by(university_id=univ_id).delete()
        
        courses_dict = univ_details.get("fees", {}).get("courses", {})
        syllabus_dict = univ_details.get("syllabus", {})
        for c_name, c_fee in courses_dict.items():
            c_syll = syllabus_dict.get(c_name, {})
            db.session.add(Course(
                university_id=univ_id,
                name=c_name,
                fee=c_fee,
                syllabus_json=json.dumps(c_syll)
            ))
            
        for f_info in univ_details.get('faculty', []):
            db.session.add(Faculty(
                university_id=univ_id,
                name=f_info.get('name'),
                designation=f_info.get('designation'),
                department=f_info.get('department'),
                email=f_info.get('email')
            ))
            
        for a_info in univ_details.get('announcements', []):
            db.session.add(Announcement(
                university_id=univ_id,
                title=a_info.get('title'),
                type=a_info.get('type', 'General'),
                desc=a_info.get('desc')
            ))
            
        for img_url in univ_details.get('gallery', []):
            db.session.add(GalleryItem(
                university_id=univ_id,
                image_url=img_url
            ))
            
        placements = univ_details.get('placements', {})
        if placements:
            db.session.add(PlacementRecord(
                university_id=univ_id,
                highest_package=placements.get('highest_package'),
                average_package=placements.get('average_package'),
                placement_rate=placements.get('placement_rate'),
                top_recruiters_json=json.dumps(placements.get('top_recruiters', []))
            ))
            
        for faq_info in univ_details.get('faqs', []):
            db.session.add(FAQItem(
                university_id=univ_id,
                question=faq_info.get('question'),
                answer=faq_info.get('answer')
            ))
        db.session.commit()
        return jsonify({'success': True, 'message': f'University {univ_id} saved successfully!'})
        
    elif request.method == 'DELETE':
        univ_id = request.args.get('id')
        if not univ_id:
            return jsonify({'success': False, 'message': 'University ID is required.'}), 400
            
        # Remove from universe_data.json
        data = load_universe_data()
        if univ_id in data:
            del data[univ_id]
            save_universe_data(data)
            
        univ = University.query.get(univ_id)
        if univ:
            Course.query.filter_by(university_id=univ_id).delete()
            Faculty.query.filter_by(university_id=univ_id).delete()
            Announcement.query.filter_by(university_id=univ_id).delete()
            GalleryItem.query.filter_by(university_id=univ_id).delete()
            PlacementRecord.query.filter_by(university_id=univ_id).delete()
            FAQItem.query.filter_by(university_id=univ_id).delete()
            Department.query.filter_by(university_id=univ_id).delete()
            Scholarship.query.filter_by(university_id=univ_id).delete()
            DocumentChunk.query.filter_by(university_id=univ_id).delete()
            
            db.session.delete(univ)
            db.session.commit()
            return jsonify({'success': True, 'message': f'University {univ_id} deleted successfully.'})
        return jsonify({'success': False, 'message': 'University not found.'}), 404

def sync_universe_json_to_database(univ_id):
    """Sync manual configuration updates from universe_data.json to SQLite database and RAG index."""
    data = load_universe_data()
    u_info = data.get(univ_id)
    if not u_info:
        return False
        
    try:
        univ = University.query.get(univ_id)
        if not univ:
            univ = University(id=univ_id, university_name=u_info.get("university_name", univ_id.upper()))
            db.session.add(univ)
            db.session.flush()
            
        # Update university fields
        univ.university_name = u_info.get("university_name", univ.university_name)
        univ.description = u_info.get("description", univ.description)
        contact = u_info.get("contact", {})
        univ.email = contact.get("email")
        univ.phone = contact.get("phone")
        univ.address = contact.get("address")
        univ.office_hours = contact.get("office_hours")
        univ.website = contact.get("website")
        univ.ranking = u_info.get("ranking")
        univ.accreditation = u_info.get("accreditation")
        univ.logo = u_info.get("logo")
        univ.brochure_url = u_info.get("brochure_url")
        univ.virtual_tour_url = u_info.get("virtual_tour_url")
        
        # Delete old child tables
        Course.query.filter_by(university_id=univ_id).delete()
        Faculty.query.filter_by(university_id=univ_id).delete()
        Announcement.query.filter_by(university_id=univ_id).delete()
        GalleryItem.query.filter_by(university_id=univ_id).delete()
        PlacementRecord.query.filter_by(university_id=univ_id).delete()
        FAQItem.query.filter_by(university_id=univ_id).delete()
        Scholarship.query.filter_by(university_id=univ_id).delete()
        
        # Create Departments if empty
        if not Department.query.filter_by(university_id=univ_id).first():
            dept_cse = Department(university_id=univ_id, name="Computer Science & Engineering", description="Department of CSE")
            db.session.add(dept_cse)
            
        # Create Courses
        courses_data = u_info.get("fees", {}).get("courses", [])
        syllabus_dict = u_info.get("syllabus", {})
        if isinstance(courses_data, list):
            for c_info in courses_data:
                c_name = c_info.get("course_name") or c_info.get("name")
                c_fee = c_info.get("tuition_fee_per_year") or c_info.get("fee")
                c_duration = c_info.get("duration") or "4 Years"
                c_syll = syllabus_dict.get(c_name, {}) if isinstance(syllabus_dict, dict) else {}
                db.session.add(Course(
                    university_id=univ_id,
                    name=c_name,
                    duration=c_duration,
                    fee=c_fee,
                    syllabus_json=json.dumps(c_syll)
                ))
        elif isinstance(courses_data, dict):
            for c_name, c_fee in courses_data.items():
                c_syll = syllabus_dict.get(c_name, {}) if isinstance(syllabus_dict, dict) else {}
                db.session.add(Course(
                    university_id=univ_id,
                    name=c_name,
                    duration="4 Years",
                    fee=c_fee,
                    syllabus_json=json.dumps(c_syll)
                ))
                
        # Create Faculty
        for f_info in u_info.get("faculty", []):
            db.session.add(Faculty(
                university_id=univ_id,
                name=f_info.get("name"),
                designation=f_info.get("designation"),
                department=f_info.get("department"),
                email=f_info.get("email")
            ))
            
        # Create Announcements
        for a_info in u_info.get("announcements", []):
            db.session.add(Announcement(
                university_id=univ_id,
                title=a_info.get("title"),
                type=a_info.get("type", "General"),
                desc=a_info.get("desc")
            ))
            
        # Create Gallery Items
        for img_url in u_info.get("gallery", []):
            db.session.add(GalleryItem(
                university_id=univ_id,
                image_url=img_url,
                caption=f"Campus View of {univ.university_name}"
            ))
            
        # Create Placement Record
        placements = u_info.get("placements", {})
        if placements:
            db.session.add(PlacementRecord(
                university_id=univ_id,
                highest_package=placements.get("highest_package"),
                average_package=placements.get("average_package"),
                placement_rate=placements.get("placement_rate"),
                top_recruiters_json=json.dumps(placements.get("top_recruiters", []))
            ))
            
        # Create FAQs
        for faq_info in u_info.get("faqs", []):
            db.session.add(FAQItem(
                university_id=univ_id,
                question=faq_info.get("question"),
                answer=faq_info.get("answer")
            ))
            
        # Create Scholarships
        for s_info in u_info.get("scholarships", []):
            db.session.add(Scholarship(
                university_id=univ_id,
                title=s_info.get("title"),
                eligibility=s_info.get("eligibility"),
                amount=s_info.get("amount")
            ))
            
        db.session.commit()
        
        # Trigger RAG chunks updates
        try:
            from services.realtime_fetcher import scrape_university_details
            import threading
            threading.Thread(target=scrape_university_details, args=(univ_id,)).start()
        except Exception as e:
            print(f"Failed to trigger chunk updates: {e}")
            
        return True
    except Exception as e:
        print(f"Error syncing config to database: {e}")
        db.session.rollback()
        return False

@app.route('/api/config', methods=['GET', 'POST'])
@rate_limit(limit=30, period=60)
def config():
    """Retrieve or update university-specific configurations."""
    univ_id = request.args.get('university_id', 'parul').strip()
    data = load_universe_data()
    
    if request.method == 'GET':
        return jsonify(data.get(univ_id, {}))
        
    elif request.method == 'POST':
        # Verify JWT Token
        token = None
        auth_header = request.headers.get('Authorization')
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]
        if not token:
            token = request.cookies.get('remember_token')
            
        if not token:
            return jsonify({'success': False, 'message': 'Authorization token is missing.'}), 401
            
        user_info = verify_jwt(token)
        if not user_info or user_info.get('role') != 'Admin':
            return jsonify({'success': False, 'message': 'Access denied.'}), 403
            
        new_data = request.get_json()
        if not new_data:
            return jsonify({'success': False, 'message': 'Invalid data provided.'}), 400
            
        data[univ_id] = new_data
        success = save_universe_data(data)
        if success:
            # Sync to database and RAG index immediately
            sync_universe_json_to_database(univ_id)
            return jsonify({'success': True, 'message': 'University details synchronized successfully!'})
        return jsonify({'success': False, 'message': 'Failed to save configuration.'}), 500

@app.route('/api/key', methods=['GET', 'POST'])
@rate_limit(limit=15, period=60)
@token_required
def manage_key():
    """Check, set or remove the Gemini API key."""
    if request.method == 'GET':
        key = get_api_key()
        masked_key = f"{key[:6]}...{key[-4:]}" if len(key) > 10 else ("*" * len(key) if key else "")
        return jsonify({'has_key': bool(key), 'masked_key': masked_key})
        
    elif request.method == 'POST':
        body = request.get_json() or {}
        new_key = body.get('api_key', '').strip()
        
        success = update_env_variable('GEMINI_API_KEY', new_key)
        if success:
            return jsonify({'success': True, 'message': 'Gemini API Key updated successfully!'})
        else:
            return jsonify({'success': False, 'message': 'Failed to update Gemini API Key.'}), 500

@app.route('/api/recommend', methods=['POST'])
def recommend():
    """Suggest best courses and universities based on budget, location, branch, interest, and rank."""
    body = request.get_json() or {}
    budget = float(body.get('budget') or 9999999)
    location = body.get('location', '').strip().lower()
    branch = body.get('branch', '').strip().lower()
    interest = body.get('interest', '').strip().lower()
    rank = int(body.get('rank') or 9999999)
    
    data = load_universe_data()
    recommendations = []
    
    for u_id, u_info in data.items():
        courses = u_info.get('fees', {}).get('courses', [])
        
        # Base cutoff threshold based on NIRF ranking
        base_cutoff = 100000
        if "iit" in u_id:
            base_cutoff = 5000
        elif "bits" in u_id:
            base_cutoff = 15000
        elif "coep" in u_id or "vjti" in u_id:
            base_cutoff = 25000
        elif "parul" in u_id:
            base_cutoff = 150000
            
        if rank < 9999999 and rank > base_cutoff * 1.5:
            continue # Exceeds entry threshold
            
        for c in courses:
            course_name = c.get('course_name', '') or c.get('name', '')
            fee_str = str(c.get('tuition_fee_per_year', '0') or c.get('fee', '0')).replace('₹', '').replace(',', '').replace('INR', '').strip()
            try:
                fee_val = float(fee_str)
            except:
                fee_val = 100000
                
            if fee_val > budget:
                continue
                
            score = 0
            
            # Location
            addr = u_info.get('contact', {}).get('address', '').lower()
            if location and location in addr:
                score += 30
                
            # Branch
            if branch and branch in course_name.lower():
                score += 40
            elif branch and any(kw in course_name.lower() for kw in branch.split()):
                score += 20
                
            # Interest
            desc = u_info.get('description', '').lower()
            if interest and interest in desc:
                score += 20
            elif interest and ('placement' in interest or 'package' in interest):
                pl = u_info.get('placements', {})
                rate_str = str(pl.get('placement_rate', '0')).replace('%', '')
                try:
                    if float(rate_str) > 85:
                        score += 15
                except:
                    pass
            elif interest and ('research' in interest or 'papers' in interest):
                papers = u_info.get('stats', {}).get('research_papers', {}).get('value', 0)
                if papers > 600:
                    score += 15
                    
            # Rank
            if rank <= base_cutoff:
                score += 10
                
            recommendations.append({
                'university_id': u_id,
                'university_name': u_info.get('university_name'),
                'location': u_info.get('contact', {}).get('address'),
                'course_name': course_name,
                'fee': c.get('tuition_fee_per_year') or c.get('fee'),
                'duration': c.get('duration'),
                'ranking': u_info.get('ranking'),
                'score': score
            })
            
    recommendations.sort(key=lambda x: x['score'], reverse=True)
    return jsonify(recommendations[:10])

@app.route('/api/predict-colleges', methods=['POST'])
def predict_colleges():
    """Predict matches for universities based on JEE Rank, Category and Home State."""
    body = request.get_json() or {}
    try:
        rank_val = body.get('rank')
        if rank_val is None:
            return jsonify({'success': False, 'message': 'JEE Rank is required.'}), 400
        rank = int(rank_val)
        if rank <= 0 or rank > 10000000:
            return jsonify({'success': False, 'message': 'Invalid JEE Rank value.'}), 400
    except (ValueError, TypeError):
        return jsonify({'success': False, 'message': 'JEE Rank must be a valid integer.'}), 400
        
    category = body.get('category', 'General').strip()
    home_state = body.get('home_state', '').strip().lower()
    
    univs = University.query.all()
    results = []
    
    for u in univs:
        # Base cutoff threshold based on NIRF ranking
        base_cutoff = 100000
        if "iit" in u.id:
            base_cutoff = 5000
        elif "bits" in u.id:
            base_cutoff = 15000
        elif "coep" in u.id or "vjti" in u.id:
            base_cutoff = 25000
        elif "parul" in u.id:
            base_cutoff = 150000
            
        # Category multiplier
        category_multiplier = 1.0
        if category == 'OBC':
            category_multiplier = 1.4
        elif category in ['SC', 'ST']:
            category_multiplier = 2.5
        elif category == 'EWS':
            category_multiplier = 1.2
            
        adjusted_cutoff = base_cutoff * category_multiplier
        
        # State quota bonus
        if home_state and u.address and home_state in u.address.lower():
            adjusted_cutoff *= 1.2
            
        probability = "Low"
        if rank <= adjusted_cutoff * 0.7:
            probability = "High"
        elif rank <= adjusted_cutoff:
            probability = "Medium"
            
        if rank > adjusted_cutoff * 1.5:
            continue
            
        results.append({
            'university_id': u.id,
            'university_name': u.university_name,
            'ranking': u.ranking or 'N/A',
            'cutoff': int(adjusted_cutoff),
            'probability': probability,
            'location': u.address or 'N/A'
        })
        
    return jsonify(sorted(results, key=lambda x: x['cutoff']))

@app.route('/api/find-scholarships', methods=['POST'])
def find_scholarships():
    """Find matching scholarships based on family income, academic marks and social category."""
    body = request.get_json() or {}
    try:
        income = float(body.get('income') or 9999999)
        marks = float(body.get('marks') or 0)
        if income < 0 or marks < 0 or marks > 100:
            return jsonify({'success': False, 'message': 'Invalid income or marks range.'}), 400
    except (ValueError, TypeError):
        return jsonify({'success': False, 'message': 'Income and marks must be valid numbers.'}), 400
        
    category = body.get('category', 'General').strip()
    
    schol_query = Scholarship.query.all()
    results = []
    
    for s in schol_query:
        elig_text = (s.eligibility or '').lower()
        
        # Income limit
        income_eligible = True
        if 'income' in elig_text or 'ews' in elig_text:
            if income > 600000:
                income_eligible = False
                
        # Marks
        marks_eligible = True
        if '>' in elig_text or '%' in elig_text:
            if '90' in elig_text and marks < 90:
                marks_eligible = False
            elif '80' in elig_text and marks < 80:
                marks_eligible = False
            elif '75' in elig_text and marks < 75:
                marks_eligible = False
                
        # Category
        category_eligible = True
        if 'sc/' in elig_text or 'st/' in elig_text or 'obc' in elig_text or 'minority' in elig_text:
            if category.lower() not in elig_text:
                category_eligible = False
                
        if income_eligible and marks_eligible and category_eligible:
            univ = University.query.get(s.university_id)
            results.append({
                'id': s.id,
                'title': s.title,
                'eligibility': s.eligibility,
                'amount': s.amount,
                'university_name': univ.university_name if univ else "Government Scheme"
            })
            
    # Add general scholarships
    if len(results) < 2:
        if income <= 800000:
            results.append({
                'id': 999,
                'title': "National Merit-cum-Means Scholarship",
                'eligibility': "Family income below 8 LPA, secured >60% in Class 12.",
                'amount': "INR 50,000 per year",
                'university_name': "National Scholarship Portal"
            })
        if category in ['SC', 'ST']:
            results.append({
                'id': 998,
                'title': "Post-Matric Scholarship Scheme for SC/ST Students",
                'eligibility': "SC/ST candidates, family income below 2.5 LPA.",
                'amount': "Full tuition fee waiver",
                'university_name': "Ministry of Social Justice"
            })
            
    return jsonify(results)

@app.route('/api/universities/<university_id>/brochure')
def download_brochure(university_id):
    """Download university course brochure dynamically in text format."""
    univ = University.query.get(university_id)
    if not univ:
        return "University not found", 404
        
    courses = Course.query.filter_by(university_id=university_id).all()
    placements = PlacementRecord.query.filter_by(university_id=university_id).first()
    faculty = Faculty.query.filter_by(university_id=university_id).all()
    
    lines = []
    lines.append("="*60)
    lines.append(f"          OFFICIAL BROCHURE - {univ.university_name.upper()}")
    lines.append("="*60)
    lines.append(f"Description: {univ.description or 'Explore course syllabus and campus life.'}\n")
    lines.append(f"Accreditation: {univ.accreditation or 'N/A'}")
    lines.append(f"NIRF Ranking: {univ.ranking or 'N/A'}")
    lines.append(f"Website: {univ.website or 'N/A'}")
    lines.append(f"Contact Email: {univ.email or 'N/A'}")
    lines.append(f"Contact Phone: {univ.phone or 'N/A'}\n")
    
    lines.append("------------------------------------------------------------")
    lines.append("                       COURSES OFFERED                      ")
    lines.append("------------------------------------------------------------")
    if courses:
        for c in courses:
            lines.append(f"- {c.name} ({c.duration or '4 Years'}): {c.fee or 'N/A'}")
    else:
        lines.append("No active courses listed.")
    lines.append("")
    
    lines.append("------------------------------------------------------------")
    lines.append("                    PLACEMENTS & SALARIES                   ")
    lines.append("------------------------------------------------------------")
    if placements:
        lines.append(f"Highest Salary Package: {placements.highest_package or 'N/A'}")
        lines.append(f"Average Salary Package: {placements.average_package or 'N/A'}")
        lines.append(f"Placement success rate: {placements.placement_rate or 'N/A'}")
    else:
        lines.append("No placement statistics listed.")
    lines.append("")
    
    lines.append("------------------------------------------------------------")
    lines.append("                       FACULTY MEMBERS                      ")
    lines.append("------------------------------------------------------------")
    if faculty:
        for f in faculty:
            lines.append(f"- {f.name} ({f.designation or 'Professor'}, {f.department or 'CSE'})")
    else:
        lines.append("No faculty roster listed.")
    lines.append("")
    
    lines.append("="*60)
    lines.append("       Generated by UniVerse AI Student Portal System       ")
    lines.append("="*60)
    
    content = "\n".join(lines)
    
    response = make_response(content)
    response.headers["Content-Disposition"] = f"attachment; filename={university_id}_brochure.txt"
    response.headers["Content-Type"] = "text/plain; charset=utf-8"
    return response

@app.route('/api/compare', methods=['GET'])
@token_required
def compare_universities():
    """Retrieve detailed side-by-side metrics of selected universities."""
    ids_param = request.args.get('ids', '')
    if not ids_param:
        return jsonify([])
    ids = [i.strip() for i in ids_param.split(',') if i.strip()]
    
    comparison_data = []
    for u_id in ids:
        univ = University.query.get(u_id)
        if not univ:
            continue
        # Get placements
        placement = PlacementRecord.query.filter_by(university_id=u_id).first()
        # Get hostel
        hostel = Hostel.query.filter_by(university_id=u_id).first()
        # Get faculty list
        faculty = Faculty.query.filter_by(university_id=u_id).all()
        # Get scholarships list
        scholarships = Scholarship.query.filter_by(university_id=u_id).all()
        # Get courses list
        courses = Course.query.filter_by(university_id=u_id).all()
        # Get events count
        events_count = Event.query.filter_by(university_id=u_id).count()
        # Get gallery count
        gallery_count = GalleryItem.query.filter_by(university_id=u_id).count()
        
        # Parse extra crawled details
        extra = {}
        if univ.crawled_details_json:
            try:
                extra = json.loads(univ.crawled_details_json)
            except:
                pass
                
        # Retrieve stats from universe_data.json
        courses_count = len(courses)
        faculty_count = len(faculty)
        real_students = 0
        real_papers = 0
        if os.path.exists(UNIVERSE_FILE):
            try:
                with open(UNIVERSE_FILE, 'r', encoding='utf-8') as f_json:
                    u_data = json.load(f_json)
                    univ_stats = u_data.get(u_id, {}).get('stats', {})
                    real_students = univ_stats.get('total_students', {}).get('value', 0)
                    real_papers = univ_stats.get('research_papers', {}).get('value', 0)
            except:
                pass
        
        if not real_students:
            real_students = 8500 + (len(univ.university_name) * 123) % 15000
        if not real_papers:
            real_papers = 450 + (len(univ.university_name) * 47) % 1200

        c_students = extra.get('total_students') or real_students
        c_faculty_count = extra.get('total_faculty') or faculty_count
        rate = placement.placement_rate if placement else 'N/A'
        avg_pkg = placement.average_package if placement else 'N/A'
        highest_pkg = placement.highest_package if placement else 'N/A'

        stats = {
            'total_students': {'value': c_students or 'Information currently unavailable'},
            'courses': {'value': courses_count or 'Information currently unavailable'},
            'faculty': {'value': c_faculty_count or 'Information currently unavailable'},
            'placement_rate': {'value': rate if '%' in str(rate) else f"{rate}%" if rate != 'N/A' else 'Information currently unavailable'},
            'average_package': {'value': avg_pkg},
            'highest_package': {'value': highest_pkg},
            'research_papers': {'value': real_papers or 'Information currently unavailable'}
        }

        comparison_data.append({
            'id': univ.id,
            'university_name': univ.university_name,
            'description': univ.description or '',
            'email': univ.email or '',
            'phone': univ.phone or '',
            'address': univ.address or '',
            'website': univ.website or '',
            'ranking': univ.ranking or '',
            'accreditation': univ.accreditation or '',
            'logo': univ.logo or '',
            'brochure_url': univ.brochure_url or '',
            'virtual_tour_url': univ.virtual_tour_url or '',
            'placements': {
                'highest_package': placement.highest_package if placement else 'N/A',
                'average_package': placement.average_package if placement else 'N/A',
                'placement_rate': placement.placement_rate if placement else 'N/A',
                'top_recruiters': json.loads(placement.top_recruiters_json) if (placement and placement.top_recruiters_json) else []
            },
            'hostel': {
                'names': hostel.names if hostel else 'N/A',
                'fees': hostel.fees if hostel else 'N/A',
                'warden_contact': hostel.warden_contact if hostel else 'N/A',
                'facilities': json.loads(hostel.facilities_json) if (hostel and hostel.facilities_json) else [],
                'details': hostel.details if hostel else 'N/A'
            },
            'faculty': [{
                'name': f.name,
                'designation': f.designation or '',
                'department': f.department or '',
                'email': f.email or ''
            } for f in faculty],
            'scholarships': [{
                'title': s.title,
                'eligibility': s.eligibility or '',
                'amount': s.amount or ''
            } for s in scholarships],
            'courses': [{
                'name': c.name,
                'duration': c.duration or '',
                'fee': c.fee or ''
            } for c in courses],
            'campus': {
                'events_count': events_count,
                'gallery_count': gallery_count
            },
            'stats': stats,
            'extra_details': extra
        })
        
    return jsonify(comparison_data)

@app.route('/api/compare/summary', methods=['POST'])
@token_required
def generate_compare_summary():
    """Generate AI summary report comparing selected universities."""
    body = request.get_json() or {}
    ids = body.get('ids', [])
    if not ids:
        return jsonify({'success': False, 'message': 'No universities selected.'}), 400
        
    if app.config.get('TESTING'):
        return jsonify({
            'success': True,
            'summary': "Mocked comparative report of selected universities for unit testing."
        })
        
    # Get comparison context
    details = []
    for u_id in ids:
        univ = University.query.get(u_id)
        if not univ:
            continue
        placement = PlacementRecord.query.filter_by(university_id=u_id).first()
        hostel = Hostel.query.filter_by(university_id=u_id).first()
        courses = Course.query.filter_by(university_id=u_id).all()
        
        info = (
            f"University: {univ.university_name}\n"
            f"Ranking: {univ.ranking or 'N/A'}\n"
            f"Accreditation: {univ.accreditation or 'N/A'}\n"
            f"Highest Package: {placement.highest_package if placement else 'N/A'}\n"
            f"Average Package: {placement.average_package if placement else 'N/A'}\n"
            f"Placement Rate: {placement.placement_rate if placement else 'N/A'}\n"
            f"Hostel Fees: {hostel.fees if hostel else 'N/A'}\n"
            f"Courses Count: {len(courses)}\n"
        )
        details.append(info)
        
    comparison_context = "\n---\n".join(details)
    
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    summary_text = ""
    if api_key:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = (
                f"You are a professional university counselor. Compare these universities side-by-side:\n\n"
                f"{comparison_context}\n\n"
                f"Provide a structured Markdown report focusing on Placements, Fees, Hostel & Campus Life, and Academics. "
                f"Give pros, cons, and a final direct recommendation depending on a student's preferences."
            )
            response = model.generate_content(prompt)
            summary_text = response.text.strip()
        except Exception as e:
            print(f"Gemini compare summary generation error: {e}")
            
    if not summary_text:
        # Rules-based clean fallback
        parts = []
        parts.append(f"### Side-by-Side Dynamic Comparison Report")
        parts.append(f"Analyzing {len(ids)} selected universities:")
        for u_id in ids:
            univ = University.query.get(u_id)
            if univ:
                placement = PlacementRecord.query.filter_by(university_id=u_id).first()
                pkg = placement.highest_package if placement else 'N/A'
                parts.append(f"- **{univ.university_name}**: NIRF ranking is `{univ.ranking or 'N/A'}` with a highest salary package of `{pkg}`.")
        parts.append("\n**Pros & Cons Overview**:")
        parts.append("- Placement-focused students should prioritize the college with the highest package rating.")
        parts.append("- Value-focused applicants should check the course fee schedules compared above.")
        parts.append("\n*Generated dynamically from synced database metrics records.*")
        summary_text = "\n".join(parts)
        
    return jsonify({
        'success': True,
        'summary': summary_text
    })

@app.route('/api/bookmarks', methods=['GET'])
@token_required
def get_bookmarks():
    """Retrieve all bookmarked items for the student."""
    user_email = request.user['user_id']
    bms = Bookmark.query.filter_by(user_id=user_email).order_by(Bookmark.id.desc()).all()
    return jsonify([bm.to_dict() for bm in bms])

@app.route('/api/bookmarks', methods=['POST'])
@token_required
def add_bookmark():
    """Toggle a bookmark for any item type."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    item_type = body.get('item_type', '').strip()
    item_id = str(body.get('item_id', '')).strip()
    university_id = body.get('university_id', '').strip()
    title = body.get('title', '').strip()
    subtitle = body.get('subtitle', '').strip()
    
    # Fallback to handle legacy university/course bookmark payloads
    if not item_id or item_id == 'None':
        if body.get('university_id'):
            item_type = 'university'
            item_id = body.get('university_id')
            university_id = item_id
            title = university_id.upper()
        elif body.get('course_id'):
            item_type = 'course'
            item_id = str(body.get('course_id'))
            title = f"Course {item_id}"
            
    if not item_id or item_id == 'None':
        return jsonify({'success': False, 'message': 'item_id is required.'}), 400
        
    # Toggle logic: remove if already bookmarked
    existing = Bookmark.query.filter_by(user_id=user_email, item_type=item_type, item_id=item_id).first()
    if existing:
        db.session.delete(existing)
        db.session.commit()
        return jsonify({'success': True, 'action': 'removed', 'message': 'Bookmark removed.'})
        
    bm = Bookmark(
        user_id=user_email,
        university_id=university_id or None,
        item_type=item_type,
        item_id=item_id,
        title=title,
        subtitle=subtitle,
        timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    )
    db.session.add(bm)
    db.session.commit()
    log_activity(user_email, f"Bookmarked {item_type}: {title or item_id}")
    return jsonify({'success': True, 'action': 'added', 'data': bm.to_dict()})

@app.route('/api/bookmarks/<int:bm_id>', methods=['DELETE'])
@token_required
def delete_bookmark_by_id(bm_id):
    """Remove bookmark by ID."""
    user_email = request.user['user_id']
    bm = Bookmark.query.filter_by(id=bm_id, user_id=user_email).first()
    if bm:
        db.session.delete(bm)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Bookmark removed.'})
    return jsonify({'success': False, 'message': 'Bookmark not found.'}), 404

@app.route('/api/recently-viewed', methods=['GET'])
@token_required
def get_recently_viewed():
    """Retrieve recently viewed items list."""
    user_email = request.user['user_id']
    items = RecentlyViewedItem.query.filter_by(user_id=user_email).order_by(RecentlyViewedItem.id.desc()).limit(10).all()
    return jsonify([item.to_dict() for item in items])

@app.route('/api/recently-viewed', methods=['POST'])
@token_required
def add_recently_viewed():
    """Add a recently viewed item context block."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    item_type = body.get('item_type', '').strip()
    item_id = str(body.get('item_id', '')).strip()
    title = body.get('title', '').strip()
    subtitle = body.get('subtitle', '').strip()
    university_id = body.get('university_id', '').strip()
    
    if not item_type or not item_id:
        return jsonify({'success': False, 'message': 'Missing item_type or item_id.'}), 400
        
    # Remove older duplicates to promote recency
    RecentlyViewedItem.query.filter_by(user_id=user_email, item_type=item_type, item_id=item_id).delete()
    
    new_item = RecentlyViewedItem(
        user_id=user_email,
        item_type=item_type,
        item_id=item_id,
        title=title,
        subtitle=subtitle,
        university_id=university_id or None,
        timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    )
    db.session.add(new_item)
    db.session.commit()
    return jsonify({'success': True, 'data': new_item.to_dict()})

@app.route('/api/notifications', methods=['GET'])
@token_required
def get_notifications():
    """Fetch notifications for the logged-in student user."""
    user_email = request.user['user_id']
    
    # Auto-seed mock categories of notifications if they do not exist for the student
    existing_titles = [n.title for n in NotificationItem.query.filter_by(user_id=user_email).all()]
    
    seed_needed = False
    for title_check in ["🚨 Admission Deadline Approaching", "💼 Placement Drive: Google India recruitment", "🎓 Merit-Based Scholarship Application", "🎉 Annual Tech Fest: UniHack 2026", "📝 Mid-Term Exam Time Table Published", "🤖 AI Recommendation: High Match University"]:
        if title_check not in existing_titles:
            seed_needed = True
            break
            
    if seed_needed:
        from datetime import datetime
        notifs_to_seed = [
            ("🚨 Admission Deadline Approaching", "The last date to submit physical documents for admission selection rounds is July 28, 2026. Please complete the form tracking workflow."),
            ("💼 Placement Drive: Google India recruitment", "Google India is hosting a career placement drive on-campus. Resume registrations are open for the software development engineer role."),
            ("🎓 Merit-Based Scholarship Application", "Scholarship allocation rounds are active. High-performing students (CGPA > 9.0) can apply to claim up to 50% tuition waiver."),
            ("🎉 Annual Tech Fest: UniHack 2026", "Join the largest university hackathon of the season, scheduled for August 12-14. Registrations are open on the events calendar."),
            ("📝 Mid-Term Exam Time Table Published", "The administrative cell has published the schedule for mid-term exams. Check the portal download links for exam room allotments."),
            ("🤖 AI Recommendation: High Match University", "Based on your interest in computer science and AI research, you have a 95% match rating for the Master of Science program at Parul University.")
        ]
        for title, msg in notifs_to_seed:
            if title not in existing_titles:
                new_notif = NotificationItem(
                    user_id=user_email,
                    title=title,
                    message=msg,
                    is_read=False,
                    timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                )
                db.session.add(new_notif)
        db.session.commit()
        
    notifs = NotificationItem.query.filter_by(user_id=user_email).order_by(NotificationItem.id.desc()).all()
    return jsonify([n.to_dict() for n in notifs])

@app.route('/api/notifications/<int:id>/read', methods=['POST'])
@token_required
def mark_read(id):
    """Mark notification as read."""
    user_email = request.user['user_id']
    n = NotificationItem.query.filter_by(id=id, user_id=user_email).first()
    if n:
        n.is_read = True
        db.session.commit()
        return jsonify({'success': True})
    return jsonify({'success': False, 'message': 'Notification not found.'}), 404

@app.route('/api/notifications/<int:id>', methods=['DELETE'])
@token_required
def delete_notification(id):
    """Delete a notification item."""
    user_email = request.user['user_id']
    n = NotificationItem.query.filter_by(id=id, user_id=user_email).first()
    if n:
        db.session.delete(n)
        db.session.commit()
        return jsonify({'success': True})
    return jsonify({'success': False, 'message': 'Notification not found.'}), 404

@app.route('/api/applications', methods=['GET'])
@token_required
def get_applications():
    """Fetch student application tracker listings."""
    user_email = request.user['user_id']
    apps = Application.query.filter_by(user_id=user_email).all()
    return jsonify([a.to_dict() for a in apps])

@app.route('/api/applications', methods=['POST'])
@token_required
def add_application():
    """Submit a new application to track."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    univ_id = body.get('university_id')
    course_name = body.get('course_name')
    status = body.get('status', 'Applied')
    notes = body.get('notes', '')
    
    if not univ_id or not course_name:
        return jsonify({'success': False, 'message': 'Missing parameters.'}), 400
        
    app_item = Application(
        user_id=user_email,
        university_id=univ_id,
        course_name=course_name,
        status=status,
        applied_date=datetime.now().strftime("%Y-%m-%d"),
        notes=notes
    )
    db.session.add(app_item)
    db.session.commit()
    log_activity(user_email, f"Submitted tracker application: {course_name} at {univ_id}")
    return jsonify({'success': True, 'data': app_item.to_dict()})

@app.route('/api/applications/<int:id>', methods=['DELETE'])
@token_required
def delete_application(id):
    """Delete an application tracker entry."""
    user_email = request.user['user_id']
    app_item = Application.query.filter_by(id=id, user_id=user_email).first()
    if app_item:
        db.session.delete(app_item)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Application tracker item deleted.'})
    return jsonify({'success': False, 'message': 'Application item not found.'}), 404

@app.route('/api/searches', methods=['GET'])
@token_required
def get_searches():
    """Retrieve recent queries searched by the student."""
    user_email = request.user['user_id']
    searches = db.session.query(RecentSearch).filter_by(user_id=user_email).order_by(RecentSearch.id.desc()).limit(10).all()
    return jsonify([s.to_dict() for s in searches])

@app.route('/api/searches', methods=['POST'])
@token_required
def add_search():
    """Add a new search query to history list."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    query = body.get('query')
    if not query:
        return jsonify({'success': False, 'message': 'Missing query.'}), 400
        
    db.session.query(RecentSearch).filter_by(user_id=user_email, query=query).delete()
    s = RecentSearch(user_id=user_email, query=query, timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    db.session.add(s)
    db.session.commit()
    return jsonify({'success': True, 'data': s.to_dict()})

@app.route('/api/profile', methods=['PUT'])
@token_required
def update_profile():
    """Update profile specifications and sync with users.json file."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    fullname = body.get('fullname')
    phone = body.get('phone')
    qualification = body.get('qualification')
    state = body.get('state')
    city = body.get('city')
    preferred_course = body.get('preferred_course')
    preferred_university = body.get('preferred_university')
    
    user = User.query.filter_by(email=user_email).first()
    if not user:
        return jsonify({'success': False, 'message': 'User not found.'}), 404
        
    user.fullname = fullname or user.fullname
    user.phone = phone or user.phone
    user.qualification = qualification or user.qualification
    user.state = state or user.state
    user.city = city or user.city
    user.preferred_course = preferred_course or user.preferred_course
    user.preferred_university = preferred_university or user.preferred_university
    
    db.session.commit()
    sync_user_to_json(user)
    log_activity(user_email, "Updated profile credentials.")
    return jsonify({'success': True, 'user': user.to_dict()})

@app.route('/api/activity-history', methods=['GET'])
@token_required
def get_activity_history():
    """Compile stats and activity audit trail for profile history dashboards."""
    user_email = request.user['user_id']
    chats = ChatSession.query.filter_by(user_id=user_email).count()
    bookmarks = Bookmark.query.filter_by(user_id=user_email).count()
    fav_courses = FavoriteCourse.query.filter_by(user_id=user_email).count()
    applications = Application.query.filter_by(user_id=user_email).count()
    searches = db.session.query(RecentSearch).filter_by(user_id=user_email).count()
    
    audit_logs = []
    if os.path.exists(LOGS_FILE):
        try:
            with open(LOGS_FILE, "r", encoding="utf-8") as f:
                logs = json.load(f)
                audit_logs = [l for l in logs if l.get('user_id') == user_email][-10:]
        except:
            pass
            
    return jsonify({
        'chats_count': chats,
        'bookmarks_count': bookmarks,
        'courses_count': fav_courses,
        'applications_count': applications,
        'searches_count': searches,
        'recent_activities': audit_logs
    })

@app.route('/api/preferences', methods=['GET'])
@token_required
def get_preferences():
    """Retrieve UI preferences configuration."""
    user_email = request.user['user_id']
    pref = UserPreference.query.filter_by(user_id=user_email).first()
    if not pref:
        pref = UserPreference(user_id=user_email)
        db.session.add(pref)
        db.session.commit()
    return jsonify(pref.to_dict())

@app.route('/api/preferences', methods=['PUT'])
@token_required
def update_preferences():
    """Save UI preferences settings."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    pref = UserPreference.query.filter_by(user_id=user_email).first()
    if not pref:
        pref = UserPreference(user_id=user_email)
        db.session.add(pref)
        
    pref.dark_mode = body.get('dark_mode', pref.dark_mode)
    pref.language = body.get('language', pref.language)
    pref.notify_email = body.get('notify_email', pref.notify_email)
    pref.notify_sms = body.get('notify_sms', pref.notify_sms)
    pref.notify_general = body.get('notify_general', pref.notify_general)
    
    db.session.commit()
    return jsonify({'success': True, 'preferences': pref.to_dict()})

# --- AUTH AND ADMIN ENDPOINTS ---

def sync_user_to_json(user_obj):
    """Sync a user model instance back to users.json file."""
    try:
        users = load_users()
        found = False
        for i, u in enumerate(users):
            if u['email'].lower() == user_obj.email.lower():
                users[i] = {
                    "fullname": user_obj.fullname,
                    "email": user_obj.email,
                    "enrollment_id": user_obj.enrollment_id,
                    "phone": user_obj.phone or '',
                    "qualification": user_obj.qualification or '',
                    "country": user_obj.country or '',
                    "state": user_obj.state or '',
                    "city": user_obj.city or '',
                    "student_status": user_obj.student_status or '',
                    "preferred_course": user_obj.preferred_course or '',
                    "preferred_university": user_obj.preferred_university or '',
                    "password_hash": user_obj.password_hash,
                    "role": user_obj.role,
                    "verified": user_obj.verified,
                    "otp_code": user_obj.otp_code,
                    "status": user_obj.status or 'Active'
                }
                found = True
                break
        if not found:
            users.append({
                "fullname": user_obj.fullname,
                "email": user_obj.email,
                "enrollment_id": user_obj.enrollment_id,
                "phone": user_obj.phone or '',
                "qualification": user_obj.qualification or '',
                "country": user_obj.country or '',
                "state": user_obj.state or '',
                "city": user_obj.city or '',
                "student_status": user_obj.student_status or '',
                "preferred_course": user_obj.preferred_course or '',
                "preferred_university": user_obj.preferred_university or '',
                "password_hash": user_obj.password_hash,
                "role": user_obj.role,
                "verified": user_obj.verified,
                "otp_code": user_obj.otp_code,
                "status": user_obj.status or 'Active'
            })
        save_users(users)
    except Exception as e:
        print(f"Error syncing user to JSON: {e}")

@app.route('/api/register', methods=['POST'])
@rate_limit(limit=10, period=60)
def register():
    """Register a new user credentials securely."""
    body = request.get_json() or {}
    fullname = body.get('fullname', '').strip()
    email = body.get('email', '').strip()
    password = body.get('password', '').strip()
    role = body.get('role', 'Student').strip()
    phone = body.get('phone', '').strip()
    qualification = body.get('qualification', '').strip()
    country = body.get('country', '').strip()
    state = body.get('state', '').strip()
    city = body.get('city', '').strip()
    student_status = body.get('student_status', 'New Student').strip()
    preferred_course = body.get('preferred_course', '').strip()
    preferred_university = body.get('preferred_university', 'parul').strip()

    if not fullname or not email or not password:
        return jsonify({'success': False, 'message': 'Full name, email, and password are required.'}), 400

    import re
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    if not re.match(email_regex, email):
        return jsonify({'success': False, 'message': 'Please enter a valid email address.'}), 400

    if len(password) < 6:
        return jsonify({'success': False, 'message': 'Password must be at least 6 characters long.'}), 400

    if role not in ['Student', 'Faculty', 'Admin']:
        return jsonify({'success': False, 'message': 'Invalid user role selected.'}), 400

    existing = User.query.filter(User.email.ilike(email)).first()
    if existing:
        return jsonify({'success': False, 'message': 'Email address already registered.'}), 400

    # Secure OTP Code generation for email verification
    import random
    otp_code = f"{random.randint(100000, 999999)}"
    total_users = User.query.count()
    enroll_id = "23030" + str(1000 + total_users + 1)
    
    new_user = User(
        fullname=fullname,
        email=email,
        enrollment_id=enroll_id,
        phone=phone,
        qualification=qualification,
        country=country,
        state=state,
        city=city,
        student_status=student_status,
        preferred_course=preferred_course,
        preferred_university=preferred_university,
        password_hash=hash_password(password),
        role=role,
        verified=False,
        otp_code=otp_code,
        status='Active'
    )
    db.session.add(new_user)
    db.session.commit()
    sync_user_to_json(new_user)
    
    print(f"[SIMULATED EMAIL SYSTEM] Sent verification code {otp_code} to {email}")
    
    return jsonify({
        'success': True,
        'message': 'Registration successful! Verification code sent to email.',
        'enrollment_id': enroll_id,
        'otp_code': otp_code,
        'preferred_university': preferred_university
    })

@app.route('/api/verify-email', methods=['POST'])
def verify_email():
    """Verify user's registration OTP code."""
    body = request.get_json() or {}
    email = body.get('email', '').strip()
    code = body.get('code', '').strip()
    
    if not email or not code:
        return jsonify({'success': False, 'message': 'Email and verification code are required.'}), 400
        
    user = User.query.filter((User.email.ilike(email)) | (User.enrollment_id == email)).first()
    if user:
        if user.otp_code == code:
            user.verified = True
            db.session.commit()
            sync_user_to_json(user)
            return jsonify({'success': True, 'message': 'Email verified successfully! You can now login.'})
        return jsonify({'success': False, 'message': 'Incorrect verification code.'}), 400
            
    return jsonify({'success': False, 'message': 'User profile not found.'}), 404

@app.route('/login', methods=['POST'])
@app.route('/api/login', methods=['POST'])
@rate_limit(limit=15, period=60)
def login():
    """Authenticate credentials and return JWT."""
    body = request.get_json() or {}
    email = body.get('email', '').strip()
    password = body.get('password', '').strip()
    remember = body.get('remember', False)
    
    if not email or not password:
        return jsonify({'success': False, 'message': 'Email/Enrollment ID and Password are required.'}), 400

    matched = User.query.filter((User.email.ilike(email)) | (User.enrollment_id == email)).first()
    
    if matched:
        if matched.status != 'Active':
            return jsonify({'success': False, 'message': 'This account has been suspended.'}), 403
            
        if not check_password(matched.password_hash, password):
            return jsonify({'success': False, 'message': 'Incorrect password.'}), 401
            
        # Check email verification status
        if not matched.verified:
            return jsonify({
                'success': False, 
                'message': 'Please verify your email address before logging in.',
                'needs_verification': True,
                'email': matched.email,
                'otp_code': matched.otp_code
            }), 403
            
        # Generate short-lived JWT access token (15 mins)
        payload = {
            'user_id': matched.email,
            'fullname': matched.fullname,
            'role': matched.role,
            'preferred_university': matched.preferred_university or 'parul'
        }
        token = generate_jwt(payload, expires_in=900)
        
        # Generate cryptographically secure refresh token (7 days)
        import secrets
        import time
        from datetime import datetime
        refresh_token = secrets.token_hex(64)
        session_expiry = int(time.time()) + (7 * 86400)
        
        # Save session to database
        new_session = UserSession(
            user_id=matched.email,
            refresh_token=refresh_token,
            expires_at=session_expiry,
            created_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        db.session.add(new_session)
        db.session.commit()
        
        # Build success response setting secure cookies
        resp = jsonify({
            'success': True,
            'message': f"Welcome back, {matched.fullname}!",
            'token': token,
            'user': matched.to_dict()
        })
        csrf_token = generate_csrf_token()
        
        # CSRF cookie
        resp.set_cookie(
            'csrf_token',
            csrf_token,
            max_age=7 * 86400 if remember else 3600,
            httponly=False,
            secure=not app.debug,
            samesite='Lax'
        )
        
        # Access token cookie
        resp.set_cookie(
            'remember_token', 
            token, 
            max_age=900, 
            httponly=False, 
            secure=not app.debug,
            samesite='Lax'
        )
        
        # Refresh token cookie (HttpOnly for security)
        resp.set_cookie(
            'refresh_token',
            refresh_token,
            max_age=7 * 86400 if remember else None,
            httponly=True,
            secure=not app.debug,
            samesite='Lax'
        )
        return resp
    else:
        return jsonify({'success': False, 'message': 'Invalid email/enrollment ID or password.'}), 401

@app.route('/api/refresh', methods=['POST'])
def refresh_token_route():
    """Generate new access token using a valid refresh token."""
    ref_token = request.cookies.get('refresh_token')
    if not ref_token:
        body = request.get_json() or {}
        ref_token = body.get('refresh_token')
        
    if not ref_token:
        return jsonify({'success': False, 'message': 'Refresh token is missing.'}), 400
        
    import time
    session_rec = UserSession.query.filter_by(refresh_token=ref_token).first()
    if not session_rec or session_rec.expires_at < int(time.time()):
        if session_rec:
            db.session.delete(session_rec)
            db.session.commit()
        return jsonify({'success': False, 'message': 'Session expired. Please log in again.'}), 401
        
    user = User.query.filter_by(email=session_rec.user_id).first()
    if not user:
        return jsonify({'success': False, 'message': 'User profile not found.'}), 404
        
    payload = {
        'user_id': user.email,
        'fullname': user.fullname,
        'role': user.role,
        'preferred_university': user.preferred_university or 'parul'
    }
    new_access_token = generate_jwt(payload, expires_in=900)
    
    resp = jsonify({
        'success': True,
        'token': new_access_token
    })
    resp.set_cookie(
        'remember_token',
        new_access_token,
        max_age=900,
        httponly=False,
        secure=not app.debug,
        samesite='Lax'
    )
    return resp

@app.route('/api/logout', methods=['POST'])
def logout():
    """Sign out, revoke refresh token, and clear cookies."""
    resp = jsonify({'success': True, 'message': 'Logged out successfully.'})
    ref_token = request.cookies.get('refresh_token')
    if ref_token:
        try:
            UserSession.query.filter_by(refresh_token=ref_token).delete()
            db.session.commit()
        except Exception:
            pass
    resp.delete_cookie('remember_token')
    resp.delete_cookie('refresh_token')
    return resp

@app.route('/api/forgot-password', methods=['POST'])
@rate_limit(limit=5, period=60)
def forgot_password():
    """Generates an OTP code for password resetting."""
    body = request.get_json() or {}
    email = body.get('email', '').strip()
    
    if not email:
        return jsonify({'success': False, 'message': 'Email address is required.'}), 400
        
    user = User.query.filter(User.email.ilike(email)).first()
    if user:
        import random
        reset_otp = f"{random.randint(100000, 999999)}"
        user.reset_otp = reset_otp
        db.session.commit()
        print(f"[SIMULATED EMAIL SYSTEM] Sent password reset code {reset_otp} to {email}")
        return jsonify({
            'success': True,
            'message': 'Password reset verification code sent to your email.',
            'otp_code': reset_otp
        })
    else:
        return jsonify({'success': False, 'message': 'No account found with this email address.'}), 404

@app.route('/api/reset-password', methods=['POST'])
@rate_limit(limit=5, period=60)
def reset_password():
    """Verifies reset OTP and updates user's password."""
    body = request.get_json() or {}
    email = body.get('email', '').strip()
    code = body.get('code', '').strip()
    new_password = body.get('new_password', '').strip()
    
    if not email or not code or not new_password:
        return jsonify({'success': False, 'message': 'Email, OTP code, and new password are required.'}), 400
        
    user = User.query.filter(User.email.ilike(email)).first()
    if user:
        if user.reset_otp == code:
            user.password_hash = hash_password(new_password)
            user.reset_otp = None
            db.session.commit()
            return jsonify({'success': True, 'message': 'Password updated successfully! You can now login.'})
        return jsonify({'success': False, 'message': 'Incorrect verification code.'}), 400
            
    return jsonify({'success': False, 'message': 'User profile not found.'}), 404

@app.route('/api/profile', methods=['GET', 'POST'])
@token_required
def user_profile():
    """Retrieve or update active user profile (JWT protected)."""
    user_email = request.user['user_id']
    user = User.query.filter(User.email.ilike(user_email)).first()
    
    if not user:
        return jsonify({'success': False, 'message': 'User profile not found.'}), 404
        
    if request.method == 'GET':
        return jsonify({
            'success': True,
            'user': user.to_dict()
        })
    elif request.method == 'POST':
        body = request.get_json() or {}
        user.fullname = body.get('fullname', user.fullname).strip()
        user.phone = body.get('phone', user.phone or '').strip()
        user.qualification = body.get('qualification', user.qualification or '').strip()
        user.country = body.get('country', user.country or '').strip()
        user.state = body.get('state', user.state or '').strip()
        user.city = body.get('city', user.city or '').strip()
        user.student_status = body.get('student_status', user.student_status or '')
        user.preferred_course = body.get('preferred_course', user.preferred_course or '').strip()
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'User profile details updated successfully!'})

@app.route('/api/logs', methods=['GET', 'DELETE'])
@token_required
def manage_logs():
    """Manage server logs file."""
    if request.user.get('role') != 'Admin':
        return jsonify({'success': False, 'message': 'Access denied.'}), 403
        
    if request.method == 'GET':
        logs = load_chat_logs()
        logs.reverse()
        try:
            page = int(request.args.get('page', 1))
            limit = int(request.args.get('limit', 20))
        except ValueError:
            page = 1
            limit = 20
        offset = (page - 1) * limit
        total = len(logs)
        paginated_logs = logs[offset:offset+limit]
        
        resp = jsonify(paginated_logs)
        resp.headers['X-Total-Count'] = str(total)
        resp.headers['X-Page'] = str(page)
        resp.headers['X-Limit'] = str(limit)
        return resp
    elif request.method == 'DELETE':
        success = save_chat_logs([])
        if success:
            return jsonify({'success': True, 'message': 'Logs cleared successfully!'})
        return jsonify({'success': False, 'message': 'Failed to clear logs.'}), 500

@app.route('/api/users', methods=['GET', 'DELETE'])
@token_required
def manage_users():
    """Manage registered users list."""
    if request.user.get('role') != 'Admin':
        return jsonify({'success': False, 'message': 'Access denied.'}), 403
    if request.method == 'GET':
        users = User.query.order_by(User.id.desc()).all()
        return jsonify([u.to_dict() for u in users])
    elif request.method == 'DELETE':
        body = request.get_json() or {}
        index = body.get('index')
        email = body.get('email')
        
        if email:
            user = User.query.filter_by(email=email).first()
            if user:
                db.session.delete(user)
                db.session.commit()
                return jsonify({'success': True, 'message': 'User deleted successfully!'})
            return jsonify({'success': False, 'message': 'User not found.'}), 404
        elif index is not None:
            users = User.query.order_by(User.id.desc()).all()
            try:
                user_to_del = users[int(index)]
                db.session.delete(user_to_del)
                db.session.commit()
                return jsonify({'success': True, 'message': 'User deleted successfully!'})
            except IndexError:
                return jsonify({'success': False, 'message': 'Index out of range.'}), 400
        else:
            User.query.filter(User.role != 'Admin').delete()
            db.session.commit()
            return jsonify({'success': True, 'message': 'All non-admin users cleared!'})

@app.route('/api/admin/upload', methods=['POST'])
@admin_required
def upload_image():
    """Upload an image to the static uploads directory."""
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file segment in the request.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected.'}), 400
        
    # Check extension
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'svg'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in allowed_extensions:
        return jsonify({'success': False, 'message': 'Invalid file extension. Allowed: PNG, JPG, JPEG, GIF, WEBP, SVG.'}), 400
        
    # Limit size to 5MB (5 * 1024 * 1024 bytes)
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)  # Reset pointer
    if file_size > 5 * 1024 * 1024:
        return jsonify({'success': False, 'message': 'File size exceeds the 5MB limit.'}), 400
        
    import secrets
    # Generate a unique secure filename to prevent collision
    filename = f"{secrets.token_hex(16)}.{ext}"
    upload_dir = os.path.join(app.root_path, 'static', 'uploads')
    os.makedirs(upload_dir, exist_ok=True)
    
    file_path = os.path.join(upload_dir, filename)
    file.save(file_path)
    
    url = f"/static/uploads/{filename}"
    return jsonify({
        'success': True,
        'message': 'File uploaded successfully!',
        'url': url
    })

@app.route('/api/admin/analytics', methods=['GET'])
@admin_required
def get_analytics():
    """Retrieve database metrics and category distribution details for overview."""
    total_univ = University.query.count()
    total_courses = Course.query.count()
    total_faculty = Faculty.query.count()
    total_students = User.query.filter_by(role='Student').count()
    
    # Load logs count
    logs = load_chat_logs()
    total_queries = len(logs)
    
    # Compute query distribution
    categories = {
        "Admissions": 0,
        "Tuition Fees": 0,
        "Exams/Policies": 0,
        "Syllabus": 0,
        "General": 0
    }
    for log in logs:
        q = log.get('query', '').lower()
        if any(w in q for w in ["admit", "admission", "apply", "deadline"]):
            categories["Admissions"] += 1
        elif any(w in q for w in ["fee", "cost", "tuition", "payment"]):
            categories["Tuition Fees"] += 1
        elif any(w in q for w in ["exam", "test", "schedule", "midterm"]):
            categories["Exams/Policies"] += 1
        elif any(w in q for w in ["syllabus", "curriculum", "subject"]):
            categories["Syllabus"] += 1
        else:
            categories["General"] += 1
            
    return jsonify({
        'success': True,
        'total_universities': total_univ,
        'total_courses': total_courses,
        'total_faculty': total_faculty,
        'total_students': total_students,
        'total_queries': total_queries,
        'query_distribution': categories
    })

@app.route('/api/announcements', methods=['GET', 'POST', 'DELETE'])
def manage_announcements():
    """Manage dynamic announcements registry."""
    if request.method == 'GET':
        ann = load_announcements()
        return jsonify(ann)
    else:
        # Verify JWT Admin Token
        token = None
        auth_header = request.headers.get('Authorization')
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]
        if not token:
            token = request.cookies.get('remember_token')
            
        if not token:
            return jsonify({'success': False, 'message': 'Authorization token is missing.'}), 401
            
        user_info = verify_jwt(token)
        if not user_info or user_info.get('role') != 'Admin':
            return jsonify({'success': False, 'message': 'Access denied.'}), 403
            
        if request.method == 'POST':
            body = request.get_json() or {}
            title = body.get('title', '').strip()
            desc = body.get('desc', '').strip()
            type_val = body.get('type', 'General').strip()
            image = body.get('image', '').strip()
            
            if not title or not desc:
                return jsonify({'success': False, 'message': 'Title and description required.'}), 400
                
            new_ann = {'title': title, 'desc': desc, 'type': type_val, 'image': image}
            ann = load_announcements()
            ann.insert(0, new_ann)
            save_announcements(ann)
            return jsonify({'success': True, 'message': 'Announcement published!'})
            
        elif request.method == 'DELETE':
            body = request.get_json() or {}
            index = body.get('index')
            ann = load_announcements()
            if index is not None:
                try:
                    ann.pop(int(index))
                    save_announcements(ann)
                    return jsonify({'success': True, 'message': 'Announcement dismissed.'})
                except IndexError:
                    return jsonify({'success': False, 'message': 'Index out of range.'}), 400
            return jsonify({'success': False, 'message': 'Delete index required.'}), 400

def rebuild_rag_document_chunks():
    """Build and store RAG document chunks for all universities in the database."""
    print("--- REBUILDING RAG KNOWLEDGE BASE CHUNKS ---")
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    universities = University.query.all()
    all_chunks_to_embed = []
    for u in universities:
        chunks = []
        
        # 1. Admissions
        admissions_text = (
            f"University: {u.university_name} ({u.id})\n"
            f"Admissions Overview:\n"
            f"Undergraduate eligibility: Pass in Class 12 exams with minimum marks and entrance score.\n"
            f"Undergraduate process: Register online, verify documents, lock choices, report to campus.\n"
            f"Undergraduate deadline: August 30, 2026\n"
            f"Postgraduate eligibility: Bachelor degree in relevant discipline from a recognized board.\n"
            f"Postgraduate process: Register online, verify documents, report to campus.\n"
            f"Postgraduate deadline: September 15, 2026"
        )
        chunks.append(('admissions', admissions_text))
        
        # 2. Hostel
        hostel_text = (
            f"University: {u.university_name} ({u.id})\n"
            f"Hostel & Mess Facilities:\n"
            f"Hostel options: Separate hostels for boys and girls with basic amenities, study tables, and Wi-Fi.\n"
            f"Hostel fees: INR 25,000 per year\n"
            f"Mess features: RO Drinking Water, 24/7 Security, clean mess catering wholesome vegetarian and non-vegetarian food options."
        )
        chunks.append(('hostel', hostel_text))
        
        # 3. Scholarships
        scholarships = Scholarship.query.filter_by(university_id=u.id).all()
        if scholarships:
            s_list = [f"- {s.title} (Amount: {s.amount}, Eligibility: {s.eligibility})" for s in scholarships]
            scholarships_text = (
                f"University: {u.university_name} ({u.id})\n"
                f"Scholarships & Financial Aid:\n" + "\n".join(s_list)
            )
        else:
            scholarships_text = (
                f"University: {u.university_name} ({u.id})\n"
                f"Scholarships & Financial Aid:\n"
                f"Merit based options: Up to 50% waiver on tuition fees for students scoring > 90% in Class 12 exams.\n"
                f"Need based options: Tuition fee waivers for students from economically weaker sections.\n"
                f"Scholarships deadline: August 30, 2026"
            )
        chunks.append(('scholarships', scholarships_text))
        
        # 4. Placements
        placement = PlacementRecord.query.filter_by(university_id=u.id).first()
        if placement:
            recs_list = []
            if placement.top_recruiters_json:
                try:
                    recs_list = json.loads(placement.top_recruiters_json)
                except Exception:
                    pass
            placements_text = (
                f"University: {u.university_name} ({u.id})\n"
                f"Placements Records & Hiring Statistics:\n"
                f"Placement rate: {placement.placement_rate}\n"
                f"Highest package: {placement.highest_package}\n"
                f"Average package: {placement.average_package}\n"
                f"Top Recruiting Hiring Partners: {', '.join(recs_list)}"
            )
            chunks.append(('placements', placements_text))
            
        # 5. Courses & Syllabus
        courses = Course.query.filter_by(university_id=u.id).all()
        for c in courses:
            syll_data = {}
            if c.syllabus_json:
                try:
                    syll_data = json.loads(c.syllabus_json)
                except Exception:
                    pass
            course_text = (
                f"University: {u.university_name} ({u.id})\n"
                f"Course Name: {c.name}\n"
                f"Tuition Fee: {c.fee}\n"
                f"Duration: {c.duration}\n"
                f"Curriculum Syllabus Details: {json.dumps(syll_data)}"
            )
            chunks.append(('courses', course_text))
            
        # 6. Faculty
        faculty = Faculty.query.filter_by(university_id=u.id).all()
        for f in faculty:
            faculty_text = (
                f"University: {u.university_name} ({u.id})\n"
                f"Faculty Member: {f.name}\n"
                f"Designation: {f.designation}\n"
                f"Department Area: {f.department}\n"
                f"Email Address: {f.email}"
            )
            chunks.append(('faculty', faculty_text))
            
        # 7. Announcements
        announcements = Announcement.query.filter_by(university_id=u.id).all()
        for a in announcements:
            announcement_text = (
                f"University: {u.university_name} ({u.id})\n"
                f"Announcement Title: {a.title}\n"
                f"Details description: {a.desc}"
            )
            chunks.append(('announcements', announcement_text))
            
        # 8. FAQs
        faqs = FAQItem.query.filter_by(university_id=u.id).all()
        for faq in faqs:
            faq_text = (
                f"University: {u.university_name} ({u.id})\n"
                f"Frequently Asked Question:\n"
                f"Question: {faq.question}\n"
                f"Answer: {faq.answer}"
            )
            chunks.append(('faqs', faq_text))
            
        # 9. General Facilities & Profile
        profile_text = (
            f"University Name: {u.university_name} ({u.id})\n"
            f"Description Profile: {u.description}\n"
            f"Location Address: {u.address}\n"
            f"Website Address: {u.website}\n"
            f"NIRF Ranking: {u.ranking}\n"
            f"NAAC Grade: {u.accreditation}\n"
            f"Campus Facilities & Amenities: Central Library, Advanced Computing Labs, Sports pavilion\n"
            f"Research areas and academic centers: Computational intelligence, sustainable energy, advanced materials"
        )
        chunks.append(('facilities', profile_text))
        
        # Accumulate chunks
        for cat, content in chunks:
            all_chunks_to_embed.append((u.id, cat, content))
            
    # Compute embeddings in memory first (no DB write lock is held during network requests)
    embedded_records = []
    for u_id, cat, content in all_chunks_to_embed:
        vector = get_embedding(content, api_key)
        embedded_records.append((u_id, cat, content, vector))
        
    # Write to DB in a single, quick transaction block
    try:
        DocumentChunk.query.delete()
        for u_id, cat, content, vector in embedded_records:
            db_chunk = DocumentChunk(
                university_id=u_id,
                category=cat,
                content=content,
                embedding_json=json.dumps(vector) if vector else None
            )
            db.session.add(db_chunk)
        db.session.commit()
    except Exception as commit_err:
        db.session.rollback()
        print(f"Error committing RAG chunks: {commit_err}")
        
    print("SUCCESS: Seeded RAG knowledge base chunks!")

with app.app_context():
    # Upgrade Bookmarks schema if old columns are missing
    try:
        db.create_all()
        db.session.execute(db.text("SELECT item_type FROM bookmarks LIMIT 1"))
    except Exception:
        db.session.rollback()
        try:
            db.session.execute(db.text("DROP TABLE IF EXISTS bookmarks"))
            db.session.commit()
            print("INFO: Dropped old bookmarks table to upgrade schema.")
        except Exception as e:
            print(f"Error dropping bookmarks table: {e}")
            
    db.create_all()
    try:
        db.session.execute(db.text("PRAGMA journal_mode=WAL"))
        db.session.commit()
    except Exception as e:
        print(f"Error setting WAL mode: {e}")
    
    # 1. Seed System Settings if empty
    if not SystemSetting.query.first():
        settings = [
            SystemSetting(key="theme", value="light-mode"),
            SystemSetting(key="allow_registration", value="true"),
            SystemSetting(key="chatbot_system_prompt", value="You are a helpful student assistant chatbot.")
        ]
        db.session.add_all(settings)
        db.session.commit()
        print("SUCCESS: Seeded System Settings.")
        
    # 2. Seed Universities & Relational Data if empty
    if not University.query.first():
        try:
            if os.path.exists(UNIVERSE_FILE):
                with open(UNIVERSE_FILE, 'r', encoding='utf-8') as f:
                    u_data = json.load(f)
                
                # Import mapping
                for u_id, u_info in u_data.items():
                    # Create University record
                    univ = University(
                        id=u_id,
                        university_name=u_info.get("university_name", u_id.upper()),
                        description=u_info.get("description") or "Explore course syllabus, annual fee sheets, and campus life.",
                        email=u_info.get("contact", {}).get("email"),
                        phone=u_info.get("contact", {}).get("phone"),
                        address=u_info.get("contact", {}).get("address"),
                        office_hours=u_info.get("contact", {}).get("office_hours"),
                        website=u_info.get("contact", {}).get("website"),
                        ranking=u_info.get("ranking"),
                        accreditation=u_info.get("accreditation"),
                        logo=u_info.get("logo"),
                        brochure_url=u_info.get("brochure_url"),
                        virtual_tour_url=u_info.get("virtual_tour_url")
                    )
                    db.session.add(univ)
                    db.session.flush()
                    
                    # Create Departments
                    dept_cse = Department(university_id=u_id, name="Computer Science & Engineering", description="Department of CSE")
                    dept_ece = Department(university_id=u_id, name="Electronics & Communication Engineering", description="Department of ECE")
                    dept_management = Department(university_id=u_id, name="Business Administration", description="Department of Management")
                    db.session.add_all([dept_cse, dept_ece, dept_management])
                    db.session.flush()
                    
                    # Create Courses (Syllabus)
                    courses_data = u_info.get("fees", {}).get("courses", [])
                    syllabus_dict = u_info.get("syllabus", {})
                    
                    if isinstance(courses_data, list):
                        for c_info in courses_data:
                            c_name = c_info.get("course_name") or c_info.get("name")
                            c_fee = c_info.get("tuition_fee_per_year") or c_info.get("fee")
                            c_duration = c_info.get("duration") or "4 Years"
                            
                            c_syll = syllabus_dict.get(c_name, {}) if isinstance(syllabus_dict, dict) else {}
                            syll_json = json.dumps(c_syll) if c_syll else "{}"
                            
                            course_record = Course(
                                university_id=u_id,
                                name=c_name,
                                duration=c_duration,
                                fee=c_fee,
                                syllabus_json=syll_json
                            )
                            db.session.add(course_record)
                    elif isinstance(courses_data, dict):
                        for c_name, c_fee in courses_data.items():
                            c_syll = syllabus_dict.get(c_name, {}) if isinstance(syllabus_dict, dict) else {}
                            syll_json = json.dumps(c_syll) if c_syll else "{}"
                            
                            duration = "4 Years" if "B.Tech" in c_name or "B.Sc" in c_name else "2 Years"
                            if "BCA" in c_name or "BBA" in c_name:
                                duration = "3 Years"
                                
                            course_record = Course(
                                university_id=u_id,
                                name=c_name,
                                duration=duration,
                                fee=c_fee,
                                syllabus_json=syll_json
                            )
                            db.session.add(course_record)
                        
                    # Create Faculty
                    fac_list = u_info.get("faculty", [])
                    for f_info in fac_list:
                        fac_record = Faculty(
                            university_id=u_id,
                            name=f_info.get("name"),
                            designation=f_info.get("designation"),
                            department=f_info.get("department"),
                            email=f_info.get("email")
                        )
                        db.session.add(fac_record)
                        
                    # Create Announcements
                    ann_list = u_info.get("announcements", [])
                    for a_info in ann_list:
                        ann_record = Announcement(
                            university_id=u_id,
                            title=a_info.get("title"),
                            type=a_info.get("type", "General"),
                            desc=a_info.get("desc")
                        )
                        db.session.add(ann_record)
                        
                    # Create Gallery Items
                    gal_list = u_info.get("gallery", [])
                    for img_url in gal_list:
                        gal_record = GalleryItem(
                            university_id=u_id,
                            image_url=img_url,
                            caption=f"Campus View of {univ.university_name}"
                        )
                        db.session.add(gal_record)
                        
                    # Create Placement Record
                    placements = u_info.get("placements", {})
                    if placements:
                        placement_record = PlacementRecord(
                            university_id=u_id,
                            highest_package=placements.get("highest_package"),
                            average_package=placements.get("average_package"),
                            placement_rate=placements.get("placement_rate"),
                            top_recruiters_json=json.dumps(placements.get("top_recruiters", []))
                        )
                        db.session.add(placement_record)
                        
                    # Create Scholarships
                    schol1 = Scholarship(
                        university_id=u_id,
                        title="Merit Scholarship",
                        eligibility="Students securing >90% in Class 12 exams.",
                        amount="50% Tuition Fee Waiver"
                    )
                    schol2 = Scholarship(
                        university_id=u_id,
                        title="Sports Excellence Scholarship",
                        eligibility="State/National level sports medalists.",
                        amount="25% Tuition Fee Waiver"
                    )
                    db.session.add_all([schol1, schol2])
                    
                    # Create FAQs
                    faqs_list = u_info.get("faqs", [])
                    for faq_info in faqs_list:
                        faq_record = FAQItem(
                            university_id=u_id,
                            question=faq_info.get("question"),
                            answer=faq_info.get("answer")
                        )
                        db.session.add(faq_record)

                    # Create Hostel details
                    hostel_data = u_info.get("hostel", {})
                    if hostel_data:
                        facs_list = hostel_data.get("facilities", ["Wi-Fi", "Common Mess", "Gym", "Power Backup", "Sports Ground"])
                        hostel_record = Hostel(
                            university_id=u_id,
                            names=hostel_data.get("names") or "Separate Boys & Girls Hostels",
                            fees=hostel_data.get("fees") or u_info.get("hostel_fees") or "₹75,000 - ₹1,40,000 per year",
                            warden_contact=hostel_data.get("warden_contact") or "9876543210",
                            facilities_json=json.dumps(facs_list),
                            details=hostel_data.get("details") or "AC and Non-AC hostel blocks with laundry services, study halls, and common dining halls."
                        )
                        db.session.add(hostel_record)

                    # Create Events
                    events_list = u_info.get("events", [])
                    if not events_list:
                        events_list = [
                            {"title": "Annual Cultural Fest", "date": "March 15, 2026", "desc": "Inter-university talent show and cultural gala."},
                            {"title": "TechExpo 2026", "date": "April 08, 2026", "desc": "Exhibition of top student projects and research papers."}
                        ]
                    for ev_data in events_list:
                        event_record = Event(
                            university_id=u_id,
                            title=ev_data.get("title"),
                            date=ev_data.get("date"),
                            desc=ev_data.get("desc")
                        )
                        db.session.add(event_record)
                        
                db.session.commit()
                print("SUCCESS: Seeded all Relational University Tables from JSON!")
        except Exception as e:
            print(f"Error seeding universities database: {e}")
            db.session.rollback()

    # 3. Seed Users if empty
    if not User.query.first():
        try:
            if os.path.exists(USERS_FILE):
                with open(USERS_FILE, 'r', encoding='utf-8') as f:
                    users_data = json.load(f)
                for u in users_data:
                    hashed = u.get('password_hash')
                    if not hashed or not (hashed.startswith('$2b$') or hashed.startswith('$2a$')):
                        pwd = "studentpassword"
                        if u.get('role') == 'Admin':
                            pwd = "adminpassword"
                        elif u.get('role') == 'Faculty':
                            pwd = "facultypassword"
                        hashed = hash_password(pwd)
                    db_user = User(
                        fullname=u['fullname'],
                        email=u['email'],
                        enrollment_id=u['enrollment_id'],
                        phone=u.get('phone', ''),
                        qualification=u.get('qualification', ''),
                        country=u.get('country', ''),
                        state=u.get('state', ''),
                        city=u.get('city', ''),
                        student_status=u.get('student_status', ''),
                        preferred_course=u.get('preferred_course', ''),
                        preferred_university=u.get('preferred_university', ''),
                        role=u.get('role', 'Student'),
                        verified=u.get('verified', True),
                        otp_code=u.get('otp_code', '123456'),
                        password_hash=hashed,
                        status=u.get('status', 'Active')
                    )
                    db.session.add(db_user)
                db.session.commit()
                print("SUCCESS: Seeded Users Database.")
        except Exception as e:
            print(f"Error seeding users database: {e}")
            db.session.rollback()

    # 4. Seed Bookmarks, Notifications, Chat History, etc.
    if not Bookmark.query.first():
        bm = Bookmark(user_id="student@univ.ac.in", university_id="parul", timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        db.session.add(bm)
        
    if not NotificationItem.query.first():
        notif = NotificationItem(
            user_id="student@univ.ac.in",
            title="Welcome to UniVerse!",
            message="Your secure dashboard portal is now ready for use.",
            is_read=False,
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        db.session.add(notif)
        
    if not ChatHistoryItem.query.first():
        session = ChatSession(
            user_id="student@univ.ac.in",
            university_id="parul",
            title="General Inquiry",
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        db.session.add(session)
        db.session.flush()
        
        chat_user = ChatHistoryItem(
            session_id=session.id,
            user_id="student@univ.ac.in",
            university_id="parul",
            sender="user",
            message="Hello chatbot",
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        chat_bot = ChatHistoryItem(
            session_id=session.id,
            user_id="student@univ.ac.in",
            university_id="parul",
            sender="bot",
            message="Hello! How can I assist you with Parul University today?",
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        db.session.add_all([chat_user, chat_bot])
        
    db.session.commit()
    db.session.remove()

    # Automatically generate and index RAG vector document chunks if empty
    # Moved to background scheduler thread to prevent blocking startup
    pass

# --- DECISION PLATFORM REST APIs ---

@app.route('/api/student/dashboard', methods=['GET'])
@token_required
def get_student_dashboard():
    """Aggregates student personalized dashboard metrics."""
    user_email = request.user['user_id']
    
    # 1. Saved/Bookmarked Universities
    bookmarks = Bookmark.query.filter_by(user_id=user_email).all()
    saved_unis = []
    for bm in bookmarks:
        univ = University.query.get(bm.university_id)
        if univ:
            saved_unis.append({
                'id': univ.id,
                'name': univ.university_name,
                'location': univ.location,
                'ranking': univ.ranking,
                'accreditation': univ.accreditation,
                'logo': univ.logo
            })
            
    # 2. Saved Courses
    fav_courses = FavoriteCourse.query.filter_by(user_id=user_email).all()
    saved_courses = []
    for fc in fav_courses:
        course = Course.query.get(fc.course_id)
        if course:
            saved_courses.append({
                'id': course.id,
                'name': course.name,
                'fee': course.fee,
                'duration': course.duration,
                'university_id': course.university_id
            })
            
    # 3. Application Tracker status
    apps = Application.query.filter_by(user_id=user_email).all()
    applications = [a.to_dict() for a in apps]
    
    # 4. Chat history counts
    chats_count = ChatSession.query.filter_by(user_id=user_email).count()
    
    # 5. Deadlines from active announcements
    announcements = Announcement.query.limit(5).all()
    deadlines = []
    for a in announcements:
        if 'last date' in a.desc.lower() or 'deadline' in a.desc.lower() or 'admission' in a.title.lower():
            deadlines.append({
                'title': a.title,
                'desc': a.desc,
                'university_id': a.university_id
            })
            
    # 6. Placement News
    placements = PlacementRecord.query.limit(4).all()
    placement_news = []
    for p in placements:
        placement_news.append({
            'university_id': p.university_id,
            'highest_package': p.highest_package,
            'average_package': p.average_package,
            'placement_rate': p.placement_rate
        })
        
    # 7. Recommended Scholarships
    scholarships = Scholarship.query.limit(3).all()
    rec_scholarships = [s.to_dict() for s in scholarships]
    
    # 8. Recently viewed (from recent searches list)
    searches = db.session.query(RecentSearch).filter_by(user_id=user_email).limit(3).all()
    recent_searches = [s.to_dict() for s in searches]

    return jsonify({
        'saved_universities': saved_unis,
        'saved_courses': saved_courses,
        'applications': applications,
        'chats_count': chats_count,
        'deadlines': deadlines,
        'placement_news': placement_news,
        'scholarship_recommendations': rec_scholarships,
        'recent_searches': recent_searches
    })

@app.route('/api/student/recommend-smart', methods=['POST'])
@token_required
def recommend_smart():
    """Generates smart matching scores and Gemini explanations for career alignment."""
    body = request.get_json() or {}
    
    budget = float(body.get('budget', 500000))
    pref_course = body.get('pref_course', '').strip().lower()
    pref_city = body.get('pref_city', '').strip().lower()
    pref_state = body.get('pref_state', '').strip().lower()
    rank = float(body.get('rank', 50000))
    career_goal = body.get('career_goal', 'Web Developer').strip()
    
    all_unis = University.query.all()
    recommendations = []
    
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    for u in all_unis:
        score = 70.0
        details = []
        
        courses = Course.query.filter_by(university_id=u.id).all()
        course_match = False
        course_fee = 0.0
        for c in courses:
            if pref_course in c.name.lower():
                course_match = True
                try:
                    fee_clean = c.fee.replace('₹', '').replace(',', '').strip()
                    course_fee = float(fee_clean)
                except Exception:
                    course_fee = 0.0
                break
                
        if course_match:
            score += 15
            details.append("Course offered")
            if course_fee <= budget:
                score += 10
                details.append("Within budget limits")
            else:
                score -= 10
                details.append("Exceeds budget")
        else:
            score -= 15
            details.append("Preferred course not directly cataloged")
            
        address_str = u.address or ""
        if pref_city and pref_city in address_str.lower():
            score += 10
            details.append("Preferred city location match")
        elif pref_state and pref_state in address_str.lower():
            score += 5
            details.append("Preferred state location match")
            
        p = PlacementRecord.query.filter_by(university_id=u.id).first()
        if p:
            try:
                avg_pkg = float(p.average_package.replace('LPA', '').replace('₹', '').strip())
                if avg_pkg > 10.0:
                    score += 5
                elif avg_pkg > 6.0:
                    score += 3
            except Exception:
                pass
                
        score = min(score, 100.0)
        score = max(score, 30.0)
        
        explanation = f"Matches your parameters with a calculated score of {score}%."
        if api_key:
            try:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = (
                    f"Explain in exactly 2 friendly sentences why {u.university_name} in {u.address or ''} "
                    f"is a good fit for a student aiming to become a '{career_goal}' with a budget of ₹{budget} per year "
                    f"and preferred course '{pref_course}'. Mention rankings ({u.ranking}) and placement average package if relevant."
                )
                response = model.generate_content(prompt)
                explanation = response.text.strip()
            except Exception as ex:
                print(f"Gemini explain error: {ex}")
                
        recommendations.append({
            'id': u.id,
            'name': u.university_name,
            'location': u.address or '',
            'ranking': u.ranking,
            'accreditation': u.accreditation,
            'score': round(score, 1),
            'explanation': explanation,
            'details': details,
            'logo': u.logo
        })
        
    recommendations.sort(key=lambda x: x['score'], reverse=True)
    return jsonify(recommendations)

@app.route('/api/student/compare-smart', methods=['POST'])
@token_required
def compare_smart():
    """Returns university comparison matrix and Gemini summary guidance."""
    body = request.get_json() or {}
    university_ids = body.get('university_ids', [])
    career_goal = body.get('career_goal', 'Software Engineer').strip()
    
    if not university_ids or len(university_ids) < 2:
        return jsonify({'success': False, 'message': 'Select at least 2 universities to compare.'}), 400
        
    compare_data = []
    for uid in university_ids:
        u = University.query.get(uid)
        if not u:
            continue
            
        courses = Course.query.filter_by(university_id=uid).all()
        placements = PlacementRecord.query.filter_by(university_id=uid).first()
        faculty = Faculty.query.filter_by(university_id=uid).limit(3).all()
        scholarships = Scholarship.query.filter_by(university_id=uid).all()
        
        fees_list = []
        for c in courses:
            try:
                fee_clean = c.fee.replace('₹', '').replace(',', '').strip()
                fees_list.append(float(fee_clean))
            except Exception:
                fees_list.append(0.0)
        avg_fee = sum(fees_list) / len(fees_list) if fees_list else 0.0

        compare_data.append({
            'id': u.id,
            'name': u.university_name,
            'location': u.address or '',
            'ranking': u.ranking,
            'accreditation': u.accreditation,
            'average_fee': avg_fee,
            'average_package': placements.average_package if placements else 'N/A',
            'highest_package': placements.highest_package if placements else 'N/A',
            'placement_rate': placements.placement_rate if placements else 'N/A',
            'faculty_count': len(faculty),
            'scholarships_count': len(scholarships)
        })
        
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    guide_summary = "Select matching profiles to retrieve AI comparative insights."
    if api_key and len(compare_data) >= 2:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = (
                f"Write a 150-word comparison summary advising a student whose goal is to become a '{career_goal}'.\n"
                f"Compare the following universities:\n"
                f"{json.dumps(compare_data, indent=2)}\n\n"
                f"Identify which university stands out as the better choice for this student and summarize reasons using bullet points."
            )
            response = model.generate_content(prompt)
            guide_summary = response.text.strip()
        except Exception as ex:
            print(f"Gemini compare summary error: {ex}")
            
    return jsonify({
        'comparison_matrix': compare_data,
        'ai_summary': guide_summary
    })

@app.route('/api/student/find-scholarships-smart', methods=['POST'])
@token_required
def find_scholarships_smart():
    """Retrieves eligible scholarships from database and generates Gemini tips, eligibility reasons, and links."""
    body = request.get_json() or {}
    income = float(body.get('income', 800000))
    category = body.get('category', 'General').strip()
    marks = float(body.get('marks', 80))
    state = body.get('state', '').strip()
    course = body.get('course', '').strip()
    
    all_schs = Scholarship.query.all()
    schs_data = [s.to_dict() for s in all_schs]
    
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    eligible = []
    ai_guidance = ""
    
    # Run AI analysis if key is configured
    if api_key and schs_data:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = (
                f"You are an AI Scholarship Advisor. Analyze the following list of scholarships:\n"
                f"{json.dumps(schs_data, indent=2)}\n\n"
                f"Determine which scholarships a student is eligible for based on:\n"
                f"- Family Income: ₹{income} per year\n"
                f"- Category: {category}\n"
                f"- Marks: {marks}%\n"
                f"- State: {state}\n"
                f"- Course: {course}\n\n"
                f"Format the output strictly as a JSON object containing:\n"
                f"{{\n"
                f"  \"eligible_scholarships\": [\n"
                f"    {{\n"
                f"      \"id\": 1,\n"
                f"      \"explanation\": \"Write a clear 1-sentence reason why they qualify based on their input marks, income, state, or category.\",\n"
                f"      \"apply_url\": \"Provide a valid official URL (e.g., official university scholarship link or https://scholarships.gov.in)\"\n"
                f"    }}\n"
                f"  ],\n"
                f"  \"ai_guidance\": \"Write a friendly 3-sentence application advice statement.\"\n"
                f"}}\n"
                f"Do not return markdown, html or formatting tags. Return only raw JSON."
            )
            response = model.generate_content(prompt)
            text = response.text.strip()
            if text.startswith("```json"):
                text = text.replace("```json", "", 1).rstrip("`").strip()
            elif text.startswith("```"):
                text = text.replace("```", "", 1).rstrip("`").strip()
            res_obj = json.loads(text)
            
            eligible_items = res_obj.get("eligible_scholarships", [])
            ai_guidance = res_obj.get("ai_guidance", "Double check individual document checklists before submitting applications.")
            
            for item in eligible_items:
                s_id = item.get("id")
                db_s = next((s for s in schs_data if s["id"] == s_id), None)
                if db_s:
                    db_s["explanation"] = item.get("explanation", "Satisfies general scheme conditions.")
                    db_s["apply_url"] = item.get("apply_url", "https://scholarships.gov.in")
                    eligible.append(db_s)
        except Exception as ex:
            print(f"Gemini scholarship advisor error: {ex}")
            eligible = []
            
    # Fallback to local rule engine if AI call failed or key is missing
    if not eligible:
        for s in schs_data:
            elig_text = s["eligibility"].lower()
            match = True
            
            # Simple local constraints check
            if "income" in elig_text or "ews" in elig_text:
                if income > 600000:
                    match = False
            if "sc" in elig_text or "st" in elig_text:
                if category not in ["SC", "ST", "OBC"]:
                    match = False
            if "marks" in elig_text or ">" in elig_text:
                if marks < 70:
                    match = False
                    
            if match:
                s["explanation"] = f"Eligible because your CGPA/Marks ({marks}%) and income (₹{income}/yr) are within specifications."
                s["apply_url"] = "https://scholarships.gov.in"
                eligible.append(s)
                
        if not ai_guidance:
            ai_guidance = "Qualified schemes computed locally. Keep income certificates and marksheets ready for registration."
            
    return jsonify({
        'scholarships': eligible,
        'ai_guidance': ai_guidance
    })

@app.route('/api/student/generate-roadmap', methods=['POST'])
@token_required
def generate_roadmap():
    """Generates structured week-by-week learning roadmaps using Gemini."""
    body = request.get_json() or {}
    track = body.get('track', 'AI').strip()
    
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    fallback_roadmaps = {
        'AI': {
            "skills": ["Python", "PyTorch", "Linear Algebra", "Calculus", "Probability", "Supervised Learning", "Deep Learning", "Transformers"],
            "projects": [
                {"name": "Matrix Algebra Library", "description": "Write basic linear algebra transformations from scratch using raw Python lists."},
                {"name": "Supervised Real-Estate Pricing", "description": "Build a predictive housing price regression model utilizing scikit-learn."},
                {"name": "Neural Network Digit Classifier", "description": "Implement a feedforward neural network in PyTorch to classify handwritten digits (MNIST)."},
                {"name": "Image Tagging Flask Application", "description": "Deploy a pre-trained computer vision model as a REST API endpoint using Flask and Docker."}
            ],
            "certifications": ["Google Cloud Professional Machine Learning Engineer", "AWS Certified Machine Learning - Specialty", "DeepLearning.AI TensorFlow Developer Certificate"],
            "timeline": "8 Weeks (approx. 10-15 hours/week)",
            "interview_prep": [
                {"question": "What is the difference between supervised and unsupervised learning?", "answer": "Supervised learning uses labeled training datasets (with known output targets), whereas unsupervised learning infers patterns and clusters from unlabeled inputs."},
                {"question": "What is overfitting and how do you prevent it?", "answer": "Overfitting occurs when a model learns the noise in the training data rather than the underlying pattern. You can prevent it using regularization (L1/L2), dropout, early stopping, or gathering more training data."}
            ],
            "resources": [
                {"name": "DeepLearning.AI Machine Learning Specialization", "url": "Coursera"},
                {"name": "3Blue1Brown: Neural Networks Visual Series", "url": "YouTube"},
                {"name": "Hugging Face NLP Course", "url": "Hugging Face Academy"}
            ],
            "roadmap": [
                {"week": "Week 1", "title": "Math Foundations", "topics": ["Linear Algebra", "Probability"], "projects": ["Matrix Library"], "resources": ["Khan Academy Linear Algebra"]},
                {"week": "Week 2", "title": "Python & NumPy", "topics": ["NumPy Arrays", "Vectorization"], "projects": ["Data Processing script"], "resources": ["Official NumPy Guide"]},
                {"week": "Week 3", "title": "Supervised Learning", "topics": ["Regression", "Classification"], "projects": ["House Price Predictor"], "resources": ["Scikit-Learn documentation"]},
                {"week": "Week 4", "title": "Unsupervised Learning", "topics": ["K-Means Clustering", "PCA"], "projects": ["Customer Segmentation"], "resources": ["StatQuest Clustering"]},
                {"week": "Week 5", "title": "Deep Learning Intro", "topics": ["Neural Networks", "Activation Functions"], "projects": ["Perceptron from Scratch"], "resources": ["3Blue1Brown Neural Networks"]},
                {"week": "Week 6", "title": "PyTorch Basics", "topics": ["Tensors", "Autograd"], "projects": ["MNIST Digit Classifier"], "resources": ["Official PyTorch Tutorials"]},
                {"week": "Week 7", "title": "Computer Vision / NLP", "topics": ["CNNs", "Transformers"], "projects": ["Image Tagging App"], "resources": ["Hugging Face Course"]},
                {"week": "Week 8", "title": "Model Deployment", "topics": ["Flask APIs", "Dockerizing Models"], "projects": ["Deploy Classifier to Web"], "resources": ["Heroku / Render guides"]}
            ]
        },
        'Cyber Security': {
            "skills": ["TCP/IP Networking", "Linux Administration", "Bash Scripting", "OWASP Top 10", "Cryptography", "Nmap", "Metasploit", "IDS/IPS"],
            "projects": [
                {"name": "Subnet Calculator", "description": "Build a network subnet and range calculator from scratch."},
                {"name": "Log Parser Script", "description": "Write a bash script that parses server auth logs and flags brute force attempts."},
                {"name": "Penetration Audit Report", "description": "Write a penetration testing report on the DVWA (Damn Vulnerable Web App) workspace."},
                {"name": "IDS Configuration", "description": "Configure Snort or pfSense firewall rules to detect network scanning signatures."}
            ],
            "certifications": ["CompTIA Security+", "Certified Ethical Hacker (CEH)", "Offensive Security Certified Professional (OSCP)"],
            "timeline": "8 Weeks (approx. 12-18 hours/week)",
            "interview_prep": [
                {"question": "What is the difference between symmetric and asymmetric encryption?", "answer": "Symmetric encryption uses the same key for both encryption and decryption. Asymmetric encryption uses a public key for encryption and a matching private key for decryption."},
                {"question": "How do you defend against SQL Injection?", "answer": "Use prepared statements (parameterized queries), input validation/sanitization, and stored procedures. Avoid constructing raw SQL queries through string concatenation."}
            ],
            "resources": [
                {"name": "PortSwigger Web Security Academy", "url": "PortSwigger"},
                {"name": "OverTheWire Bandit Command Line Game", "url": "OverTheWire"},
                {"name": "TCM Security: Ethical Hacking Academy", "url": "TCM Academy"}
            ],
            "roadmap": [
                {"week": "Week 1", "title": "Networking Basics", "topics": ["TCP/IP", "DNS", "Subnetting"], "projects": ["Network Subnet Calculator"], "resources": ["Professor Messer CompTIA Network+"]},
                {"week": "Week 2", "title": "Linux & Command Line", "topics": ["Bash scripting", "Permissions"], "projects": ["Log parser script"], "resources": ["OverTheWire Bandit game"]},
                {"week": "Week 3", "title": "Cyber Threat Landscape", "topics": ["Malware", "Phishing", "Social Engineering"], "projects": ["Create educational phishing guide"], "resources": ["CISA Alerts"]},
                {"week": "Week 4", "title": "Web Application Security", "topics": ["OWASP Top 10", "SQL Injection", "XSS"], "projects": ["Penetration test on DVWA"], "resources": ["PortSwigger Web Security Academy"]},
                {"week": "Week 5", "title": "Cryptography Foundations", "topics": ["Symmetric keys", "PKI", "Hashing"], "projects": ["Secure file encryption tool"], "resources": ["Coursera Cryptography"]},
                {"week": "Week 6", "title": "Network Scanning", "topics": ["Nmap", "Wireshark packeting"], "projects": ["Network vulnerability audit report"], "resources": ["Wireshark docs"]},
                {"week": "Week 7", "title": "Penetration Testing", "topics": ["Metasploit", "Privilege Escalation"], "projects": ["HackTheBox machine challenge"], "resources": ["TCM Security Practical Ethical Hacking"]},
                {"week": "Week 8", "title": "Defensive Security", "topics": ["SIEM", "Incident Analysis", "Firewalls"], "projects": ["Setup pfSense or Snort IDS"], "resources": ["Blue Team Labs"]}
            ]
        },
        'Cloud': {
            "skills": ["AWS Core Services", "Linux Shell", "Cloud Security (IAM)", "NoSQL Cloud DBs", "Serverless Functions", "Terraform IaC", "CI/CD Pipelines", "Docker"],
            "projects": [
                {"name": "VPC Deploy static site", "description": "Configure VPC subnets and host a static site on a secure EC2 instance."},
                {"name": "Secure Cloud Backup", "description": "Build an automated server backup pipeline uploading files into encrypted S3 buckets."},
                {"name": "Image Resizing Lambda", "description": "Write a Python serverless function that automatically resizes uploaded images on S3 triggers."},
                {"name": "IaC VPC Setup", "description": "Write Terraform scripts to construct a multi-tier server architecture on AWS."}
            ],
            "certifications": ["AWS Certified Cloud Practitioner", "AWS Certified Solutions Architect - Associate", "HashiCorp Certified: Terraform Associate"],
            "timeline": "8 Weeks (approx. 10-12 hours/week)",
            "interview_prep": [
                {"question": "What is infrastructure as code (IaC)?", "answer": "IaC is the management and provisioning of cloud infrastructure resources through machine-readable definition files (like Terraform or CloudFormation) rather than manual interactive configuration."},
                {"question": "Explain serverless computing.", "answer": "Serverless is a execution model where the cloud provider manages server allocation and scaling dynamically. Developers focus only on application logic (e.g., AWS Lambda) and are billed only for execution runtime."}
            ],
            "resources": [
                {"name": "AWS Cloud Practitioner Training Path", "url": "AWS Skill Builder"},
                {"name": "Terraform Certification Guide", "url": "HashiCorp Learn"},
                {"name": "AWS Serverless Developer Guide", "url": "AWS Documentation"}
            ],
            "roadmap": [
                {"week": "Week 1", "title": "Cloud Computing Essentials", "topics": ["SaaS/PaaS/IaaS", "Cloud economics"], "projects": ["Cost calculation simulation"], "resources": ["AWS Cloud Practitioner essentials"]},
                {"week": "Week 2", "title": "Core Virtualization", "topics": ["EC2", "Virtual Machines", "VPCs"], "projects": ["Deploy static site in VPC EC2"], "resources": ["AWS VPC tutorials"]},
                {"week": "Week 3", "title": "Cloud Storage Solutions", "topics": ["S3 bucket policy", "EBS volume options"], "projects": ["Secure cloud backup pipeline"], "resources": ["AWS S3 Developer Guide"]},
                {"week": "Week 4", "title": "Identity & Access Management", "topics": ["IAM roles", "MFA enforcement"], "projects": ["Setup strict RBAC cloud structure"], "resources": ["AWS Security Center"]},
                {"week": "Week 5", "title": "Cloud Databases", "topics": ["RDS", "DynamoDB NoSQL"], "projects": ["Create serverless backend database"], "resources": ["AWS RDS specs"]},
                {"week": "Week 6", "title": "Serverless Architecture", "topics": ["Lambda Functions", "API Gateway"], "projects": ["Image resizing serverless tool"], "resources": ["AWS Serverless docs"]},
                {"week": "Week 7", "title": "Infrastructure as Code", "topics": ["Terraform", "CloudFormation"], "projects": ["Terraform script for multi-tier VPC"], "resources": ["HashiCorp Learn"]},
                {"week": "Week 8", "title": "DevOps & Pipelines", "topics": ["CI/CD pipelines", "Monitoring Tools"], "projects": ["Automated GitHub Actions deploy pipeline"], "resources": ["AWS DevOps center"]}
            ]
        },
        'Data Science': {
            "skills": ["Python / Pandas", "SQL Queries", "Matplotlib / Seaborn", "Statistical Modeling", "Hypothesis Testing", "Linear Regression", "Hyperparameter Tuning", "Tableau / PowerBI"],
            "projects": [
                {"name": "Dataset Cleaner", "description": "Write Python script using Pandas to identify, clean, and resolve nulls or outliers inside dirty retail CSVs."},
                {"name": "Retail Business Dashboard", "description": "Draft complex SQL queries joining multiple tables to render business KPI charts."},
                {"name": "A/B Conversion Analysis", "description": "Use hypothesis testing to evaluate user click-through conversions for landing page variants."},
                {"name": "Survival Predictor model", "description": "Build a predictive machine learning model using logistic regression to forecast passenger survival probabilities."}
            ],
            "certifications": ["Google Data Analytics Professional Certificate", "Microsoft Certified: Power BI Data Analyst Associate", "IBM Data Science Professional Certificate"],
            "timeline": "8 Weeks (approx. 10-15 hours/week)",
            "interview_prep": [
                {"question": "What is an outlier and how do you handle it?", "answer": "An outlier is a data point that differs significantly from other observations. You can handle them by removing them (if they are measurement errors), transforming the variable (e.g., log transformation), or using robust modeling algorithms."},
                {"question": "Explain A/B testing.", "answer": "A/B testing is a statistical methodology comparing two versions (A and B) of a variable (e.g., web pages) to determine which performs better based on statistical significance."}
            ],
            "resources": [
                {"name": "SQLZoo Database Sandbox Tutorials", "url": "SQLZoo"},
                {"name": "Kaggle Data Science Guides", "url": "Kaggle Learn"},
                {"name": "OpenIntro Statistics Textbook", "url": "OpenIntro"}
            ],
            "roadmap": [
                {"week": "Week 1", "title": "Python Basics & Pandas", "topics": ["Pandas DataFrames", "Data Wrangling"], "projects": ["Clean dirty CSV dataset"], "resources": ["Kaggle Data Cleaning Course"]},
                {"week": "Week 2", "title": "SQL for Analytics", "topics": ["SQL Joins", "Aggregations"], "projects": ["Mock retail business analytics dashboard"], "resources": ["SQLZoo"]},
                {"week": "Week 3", "title": "Data Visualization", "topics": ["Matplotlib", "Seaborn"], "projects": ["Visualize Covid or Stock Trends"], "resources": ["Storytelling with Data"]},
                {"week": "Week 4", "title": "Exploratory Data Analysis", "topics": ["Outliers", "Correlation Analysis"], "projects": ["Full EDA Report on housing datasets"], "resources": ["TowardsDataScience EDA"]},
                {"week": "Week 5", "title": "Statistical Inference", "topics": ["Hypothesis Testing", "A/B Testing"], "projects": ["A/B Testing analysis for conversions"], "resources": ["OpenIntro Statistics"]},
                {"week": "Week 6", "title": "Predictive Modeling", "topics": ["Linear Regression", "Logistic Regression"], "projects": ["Titanic survival predictor"], "resources": ["Kaggle Titanic"]},
                {"week": "Week 7", "title": "Machine Learning Workflows", "topics": ["Cross Validation", "Grid Search"], "projects": ["Hyperparameter tuning project"], "resources": ["Machine Learning Mastery"]},
                {"week": "Week 8", "title": "PowerBI / Tableau", "topics": ["BI Dashboards", "KPI reporting"], "projects": ["Build corporate sales dashboard"], "resources": ["Tableau Public Gallery"]}
            ]
        },
        'Web Development': {
            "skills": ["HTML5 / CSS Layouts", "JavaScript DOM", "React Frontend", "Express / Node.js", "REST APIs", "PostgreSQL", "MongoDB", "Deployment"],
            "projects": [
                {"name": "Responsive Landing Page", "description": "Construct a responsive tech agency landing page using Flexbox and Grid."},
                {"name": "Weather Dashboard Web App", "description": "Build a React weather forecasting interface calling live geographical weather APIs."},
                {"name": "Note-Taking API Service", "description": "Write a Node.js REST API with full CRUD endpoints using Express."},
                {"name": "Full Stack MERN/PERN App", "description": "Develop and deploy a full stack web app utilizing user authentication and profile state database live on Render/Vercel."}
            ],
            "certifications": ["Meta Front-End Developer Professional Certificate", "FreeCodeCamp Responsive Web Design Certificate", "MongoDB Certified Developer Associate"],
            "timeline": "8 Weeks (approx. 10-15 hours/week)",
            "interview_prep": [
                {"question": "What is the difference between state and props in React?", "answer": "State is a local data storage that is managed internally by the component itself. Props are external configurations passed down from parent to child components and are read-only."},
                {"question": "What is a REST API?", "answer": "A REST API is an application programming interface adhering to REST constraints, utilizing HTTP requests (GET, POST, PUT, DELETE) to fetch, create, update, or delete server-side data."}
            ],
            "resources": [
                {"name": "MDN Web Docs", "url": "MDN Web"},
                {"name": "Javascript Info Guides", "url": "JavaScript.info"},
                {"name": "Full Stack Open Course", "url": "University of Helsinki"}
            ],
            "roadmap": [
                {"week": "Week 1", "title": "HTML5 & CSS3 layouts", "topics": ["Flexbox", "Grid", "Responsiveness"], "projects": ["Responsive Portfolio Landing Page"], "resources": ["MDN Web Docs HTML/CSS"]},
                {"week": "Week 2", "title": "JavaScript DOM", "topics": ["DOM manipulation", "ES6 Features"], "projects": ["Dynamic interactive task board"], "resources": ["JavaScript Info"]},
                {"week": "Week 3", "title": "Modern Frontend (React)", "topics": ["React Components", "Hooks (useState)"], "projects": ["Weather Dashboard App"], "resources": ["Official React Docs"]},
                {"week": "Week 4", "title": "State Management & Routing", "topics": ["React Context", "React Router"], "projects": ["Mini e-commerce checkout page"], "resources": ["Redux Toolkit tutorials"]},
                {"week": "Week 5", "title": "Backend Development", "topics": ["Node.js / Express", "REST APIs"], "projects": ["Note taking API service"], "resources": ["Express.js Guide"]},
                {"week": "Week 6", "title": "Relational Databases", "topics": ["PostgreSQL / SQL basics"], "projects": ["User Authentication and Profile database"], "resources": ["PostgreSQL Tutorial"]},
                {"week": "Week 7", "title": "NoSQL Options", "topics": ["MongoDB", "Mongoose schema"], "projects": ["Social Media feed backend API"], "resources": ["MongoDB University"]},
                {"week": "Week 8", "title": "Production Deployment", "topics": ["Vercel", "Render", "CI/CD checks"], "projects": ["Deploy full MERN / PERN site live"], "resources": ["Render deploy docs"]}
            ]
        }
    }
    
    fallback = fallback_roadmaps.get(track, fallback_roadmaps['AI'])
    
    if api_key:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = (
                f"Generate a detailed learning path for a student aiming to learn '{track}'.\n"
                f"Format the output strictly as a JSON object with the following structure:\n"
                f"{{\n"
                f"  \"skills\": [\"Skill A\", \"Skill B\"],\n"
                f"  \"projects\": [\n"
                f"    {{\"name\": \"Project Name\", \"description\": \"Short description\"}}\n"
                f"  ],\n"
                f"  \"certifications\": [\"Cert A\", \"Cert B\"],\n"
                f"  \"timeline\": \"Overall timeline description (e.g. 8 Weeks, 10 hours/week)\",\n"
                f"  \"interview_prep\": [\n"
                f"    {{\"question\": \"Question text\", \"answer\": \"Answer text\"}}\n"
                f"  ],\n"
                f"  \"resources\": [\n"
                f"    {{\"name\": \"Resource Name\", \"url\": \"URL or Platform (e.g. Coursera, MDN)\"}}\n"
                f"  ],\n"
                f"  \"roadmap\": [\n"
                f"    {{\n"
                f"      \"week\": \"Week 1\",\n"
                f"      \"title\": \"Week Title\",\n"
                f"      \"topics\": [\"Topic A\", \"Topic B\"],\n"
                f"      \"projects\": [\"Small Project Name\"],\n"
                f"      \"resources\": [\"Resource Name\"]\n"
                f"    }}\n"
                f"  ]\n"
                f"}}\n\n"
                f"Do not return any other text, html wrappers, or markdown backticks, only valid raw JSON."
            )
            response = model.generate_content(prompt)
            text = response.text.strip()
            if text.startswith("```json"):
                text = text.replace("```json", "", 1).rstrip("`").strip()
            elif text.startswith("```"):
                text = text.replace("```", "", 1).rstrip("`").strip()
                
            parsed = json.loads(text)
            return jsonify(parsed)
        except Exception as e:
            app.logger.warning(f"Roadmap generation failed with Gemini, using fallback: {e}")
            return jsonify(fallback)
    else:
        return jsonify(fallback)

@app.route('/api/student/generate-resume', methods=['POST'])
@token_required
def generate_resume():
    """Generates optimized bullet points for student resume draft."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    
    fullname = body.get('fullname', '').strip()
    email = body.get('email', '').strip()
    phone = body.get('phone', '').strip()
    summary = body.get('summary', '').strip()
    skills = body.get('skills', [])
    education = body.get('education', [])
    experience = body.get('experience', [])
    projects = body.get('projects', [])
    
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    optimized_summary = summary
    optimized_projects = projects
    optimized_experience = experience
    
    if api_key and (summary or projects or experience):
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            if summary:
                prompt_sum = f"Rewrite this resume professional summary in 2 sentences to make it sound highly professional and ATS-friendly: '{summary}'"
                resp_sum = model.generate_content(prompt_sum)
                optimized_summary = resp_sum.text.strip()
                
            if projects:
                for idx, p in enumerate(projects):
                    desc = p.get('desc', '')
                    if desc:
                        prompt_p = f"Rewrite this project description bullet list using action verbs and technical keywords: '{desc}'"
                        resp_p = model.generate_content(prompt_p)
                        optimized_projects[idx]['desc'] = resp_p.text.strip()
        except Exception as ex:
            print(f"Gemini resume optimization error: {ex}")
            
    p_row = ResumeProfile.query.filter_by(user_id=user_email).first()
    if not p_row:
        p_row = ResumeProfile(user_id=user_email, fullname=fullname, email=email, phone=phone, timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        db.session.add(p_row)
        
    p_row.fullname = fullname
    p_row.email = email
    p_row.phone = phone
    p_row.summary = optimized_summary
    p_row.skills = json.dumps(skills)
    p_row.education = json.dumps(education)
    p_row.experience = json.dumps(optimized_experience)
    p_row.projects = json.dumps(optimized_projects)
    db.session.commit()
    
    return jsonify(p_row.to_dict())

@app.route('/api/student/download-resume-pdf')
@token_required
def download_resume_pdf():
    """Generates an ATS-friendly, professionally styled PDF from Resume Profile data."""
    user_email = request.user['user_id']
    profile = ResumeProfile.query.filter_by(user_id=user_email).first()
    if not profile:
        return jsonify({'success': False, 'message': 'Resume profile not found. Please create one first.'}), 404
        
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter, 
        rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
    )
    
    styles = getSampleStyleSheet()
    
    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#0F172A'),
        alignment=1
    )
    
    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#475569'),
        alignment=1
    )
    
    section_title_style = ParagraphStyle(
        'SectionTitleStyle',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=3
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=colors.HexColor('#334155')
    )
    
    story = []
    
    # Name & Contact
    story.append(Paragraph(profile.fullname or 'Anonymous Student', name_style))
    story.append(Spacer(1, 4))
    contact_text = f"{profile.email or ''}  |  {profile.phone or ''}"
    story.append(Paragraph(contact_text, contact_style))
    story.append(Spacer(1, 12))
    
    def add_section(title_text):
        story.append(Paragraph(title_text, section_title_style))
        divider = Table([['']], colWidths=[532], rowHeights=[1.5])
        divider.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#CBD5E1')),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(divider)
        story.append(Spacer(1, 6))

    # Summary
    if profile.summary:
        add_section("PROFESSIONAL SUMMARY")
        story.append(Paragraph(profile.summary, body_style))
        story.append(Spacer(1, 10))
        
    # Skills
    try:
        skills_list = json.loads(profile.skills or '[]')
    except:
        skills_list = []
    if skills_list:
        add_section("CORE TECHNICAL SKILLS")
        skills_str = ", ".join(skills_list)
        story.append(Paragraph(skills_str, body_style))
        story.append(Spacer(1, 10))
        
    # Experience
    try:
        exp_list = json.loads(profile.experience or '[]')
    except:
        exp_list = []
    if exp_list:
        add_section("PROFESSIONAL EXPERIENCE")
        for exp in exp_list:
            role = exp.get('role', '')
            company = exp.get('company', '')
            duration = exp.get('duration', '')
            desc = exp.get('desc', '')
            
            hdr_text = f"<b>{role}</b>  -  {company}"
            story.append(Paragraph(hdr_text, body_style))
            story.append(Paragraph(f"<font color='#64748B'><i>{duration}</i></font>", body_style))
            if desc:
                bullets = [b.strip() for b in desc.split('\n') if b.strip()]
                for b in bullets:
                    bullet_text = f"• {b}"
                    story.append(Paragraph(bullet_text, body_style))
            story.append(Spacer(1, 8))
            
    # Projects
    try:
        proj_list = json.loads(profile.projects or '[]')
    except:
        proj_list = []
    if proj_list:
        add_section("KEY PROJECTS PORTFOLIO")
        for proj in proj_list:
            name = proj.get('name', '')
            tech = proj.get('tech', '')
            desc = proj.get('desc', '')
            
            hdr_text = f"<b>{name}</b>  -  <font color='#475569'><i>{tech}</i></font>"
            story.append(Paragraph(hdr_text, body_style))
            if desc:
                bullets = [b.strip() for b in desc.split('\n') if b.strip()]
                for b in bullets:
                    bullet_text = f"• {b}"
                    story.append(Paragraph(bullet_text, body_style))
            story.append(Spacer(1, 8))
            
    # Education
    try:
        edu_list = json.loads(profile.education or '[]')
    except:
        edu_list = []
    if edu_list:
        add_section("EDUCATION")
        for edu in edu_list:
            degree = edu.get('degree', '')
            school = edu.get('school', '')
            year = edu.get('year', '')
            
            hdr_text = f"<b>{degree}</b>  -  {school}  ({year})"
            story.append(Paragraph(hdr_text, body_style))
            story.append(Spacer(1, 4))
            
    doc.build(story)
    buffer.seek(0)
    
    from flask import send_file
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{profile.fullname.replace(' ', '_')}_Resume.pdf",
        mimetype='application/pdf'
    )

@app.route('/api/student/download-resume-docx')
@token_required
def download_resume_docx():
    """Generates an ATS-friendly, clean MS Word DOCX from Resume Profile data."""
    user_email = request.user['user_id']
    profile = ResumeProfile.query.filter_by(user_id=user_email).first()
    if not profile:
        return jsonify({'success': False, 'message': 'Resume profile not found. Please create one first.'}), 404
        
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    def set_font(run, name="Calibri", size=10, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            from docx.shared import RGBColor
            run.font.color.rgb = color
            
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = name_p.add_run(profile.fullname or 'Anonymous Student')
    set_font(r_name, "Calibri", 18, bold=True)
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{profile.email or ''}  |  {profile.phone or ''}"
    r_contact = contact_p.add_run(contact_text)
    set_font(r_contact, "Calibri", 10)
    
    def add_section_header(title):
        h_p = doc.add_paragraph()
        h_p.paragraph_format.space_before = Pt(10)
        h_p.paragraph_format.space_after = Pt(2)
        h_run = h_p.add_run(title)
        set_font(h_run, "Calibri", 11, bold=True)
        
        border_p = doc.add_paragraph()
        border_p.paragraph_format.space_before = Pt(0)
        border_p.paragraph_format.space_after = Pt(4)
        b_run = border_p.add_run("―" * 60)
        set_font(b_run, "Calibri", 8, bold=True)
        
    if profile.summary:
        add_section_header("PROFESSIONAL SUMMARY")
        sum_p = doc.add_paragraph()
        sum_p.paragraph_format.space_after = Pt(8)
        r_sum = sum_p.add_run(profile.summary)
        set_font(r_sum, "Calibri", 10)
        
    try:
        skills_list = json.loads(profile.skills or '[]')
    except:
        skills_list = []
    if skills_list:
        add_section_header("CORE TECHNICAL SKILLS")
        skills_p = doc.add_paragraph()
        skills_p.paragraph_format.space_after = Pt(8)
        r_skills = skills_p.add_run(", ".join(skills_list))
        set_font(r_skills, "Calibri", 10)
        
    try:
        exp_list = json.loads(profile.experience or '[]')
    except:
        exp_list = []
    if exp_list:
        add_section_header("PROFESSIONAL EXPERIENCE")
        for exp in exp_list:
            role = exp.get('role', '')
            company = exp.get('company', '')
            duration = exp.get('duration', '')
            desc = exp.get('desc', '')
            
            exp_p = doc.add_paragraph()
            exp_p.paragraph_format.space_before = Pt(4)
            exp_p.paragraph_format.space_after = Pt(0)
            r_title = exp_p.add_run(f"{role}  |  {company}")
            set_font(r_title, "Calibri", 10, bold=True)
            
            dur_p = doc.add_paragraph()
            dur_p.paragraph_format.space_after = Pt(2)
            r_dur = dur_p.add_run(duration)
            set_font(r_dur, "Calibri", 9, italic=True)
            
            if desc:
                bullets = [b.strip() for b in desc.split('\n') if b.strip()]
                for b in bullets:
                    bullet_p = doc.add_paragraph(style='List Bullet')
                    bullet_p.paragraph_format.space_after = Pt(2)
                    r_b = bullet_p.add_run(b)
                    set_font(r_b, "Calibri", 10)
                    
    try:
        proj_list = json.loads(profile.projects or '[]')
    except:
        proj_list = []
    if proj_list:
        add_section_header("KEY PROJECTS PORTFOLIO")
        for proj in proj_list:
            name = proj.get('name', '')
            tech = proj.get('tech', '')
            desc = proj.get('desc', '')
            
            proj_p = doc.add_paragraph()
            proj_p.paragraph_format.space_before = Pt(4)
            proj_p.paragraph_format.space_after = Pt(2)
            r_title = proj_p.add_run(f"{name}  -  ({tech})")
            set_font(r_title, "Calibri", 10, bold=True)
            
            if desc:
                bullets = [b.strip() for b in desc.split('\n') if b.strip()]
                for b in bullets:
                    bullet_p = doc.add_paragraph(style='List Bullet')
                    bullet_p.paragraph_format.space_after = Pt(2)
                    r_b = bullet_p.add_run(b)
                    set_font(r_b, "Calibri", 10)
                    
    try:
        edu_list = json.loads(profile.education or '[]')
    except:
        edu_list = []
    if edu_list:
        add_section_header("EDUCATION")
        for edu in edu_list:
            degree = edu.get('degree', '')
            school = edu.get('school', '')
            year = edu.get('year', '')
            
            edu_p = doc.add_paragraph()
            edu_p.paragraph_format.space_after = Pt(4)
            r_edu = edu_p.add_run(f"{degree}  -  {school}  ({year})")
            set_font(r_edu, "Calibri", 10)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    from flask import send_file
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{profile.fullname.replace(' ', '_')}_Resume.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/api/student/admission/checklist', methods=['GET'])
@token_required
def admission_checklist():
    """Generates admission checklist for selected university and fetches completed tasks."""
    user_email = request.user['user_id']
    univ_id = request.args.get('university_id', '').strip()
    if not univ_id:
        return jsonify({'success': False, 'message': 'University ID is required.'}), 400
        
    univ = University.query.get(univ_id)
    if not univ:
        return jsonify({'success': False, 'message': 'University not found.'}), 404
        
    # Get database checked items list
    row = AdmissionChecklist.query.filter_by(user_id=user_email, university_id=univ_id).first()
    completed_steps = json.loads(row.completed_steps) if row else []
    
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    checklist_data = None
    if api_key:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = (
                f"You are an AI Admission Advisor. Generate a structured admission checklist for the following university:\n"
                f"Name: {univ.university_name}\n"
                f"Description: {univ.description}\n"
                f"Extra Info: {univ.crawled_details_json or '{}'}\n\n"
                f"Generate details for:\n"
                f"1. Required Documents\n"
                f"2. Application Steps\n"
                f"3. Important Dates\n"
                f"4. Fees\n"
                f"5. Hostel Process\n"
                f"6. Scholarship Process\n\n"
                f"Format the output strictly as a JSON object containing:\n"
                f"{{\n"
                f"  \"documents\": [{{ \"key\": \"doc_1\", \"label\": \"Required document detail...\" }}],\n"
                f"  \"steps\": [{{ \"key\": \"step_1\", \"label\": \"Application step detail...\" }}],\n"
                f"  \"dates\": [{{ \"key\": \"date_1\", \"label\": \"Date detail...\" }}],\n"
                f"  \"fees\": [{{ \"key\": \"fee_1\", \"label\": \"Fee detail...\" }}],\n"
                f"  \"hostel\": [{{ \"key\": \"hostel_1\", \"label\": \"Hostel booking step...\" }}],\n"
                f"  \"scholarship\": [{{ \"key\": \"scholar_1\", \"label\": \"Scholarship step...\" }}]\n"
                f"}}\n"
                f"Provide at least 3-4 items for each list. Each key must be unique (e.g. doc_1, doc_2, step_1, step_2, etc.).\n"
                f"Do not return markdown, html, or formatting tags. Return only raw JSON."
            )
            response = model.generate_content(prompt)
            text = response.text.strip()
            if text.startswith("```json"):
                text = text.replace("```json", "", 1).rstrip("`").strip()
            elif text.startswith("```"):
                text = text.replace("```", "", 1).rstrip("`").strip()
            checklist_data = json.loads(text)
        except Exception as ex:
            print(f"Gemini admission checklist generation error: {ex}")
            
    if not checklist_data:
        # Fallback pre-seeded lists
        checklist_data = {
            "documents": [
                { "key": "doc_1", "label": "Class 10th & 12th Marksheets (Original & copies)" },
                { "key": "doc_2", "label": "School Leaving Certificate / Transfer Certificate" },
                { "key": "doc_3", "label": "Migration Certificate (for other state boards)" },
                { "key": "doc_4", "label": "Passport-sized photographs (6 copies)" },
                { "key": "doc_5", "label": "Income & Caste Verification certificate (if applicable)" }
            ],
            "steps": [
                { "key": "step_1", "label": f"Register on the official {univ.university_name} portal" },
                { "key": "step_2", "label": "Fill program preferences and academic details" },
                { "key": "step_3", "label": "Upload soft copies of marks cards and certificates" },
                { "key": "step_4", "label": "Pay the admission application form fee online" },
                { "key": "step_5", "label": "Submit application form and keep printed copy ready" }
            ],
            "dates": [
                { "key": "date_1", "label": "Application Portal Opens: June 15, 2026" },
                { "key": "date_2", "label": "Last Date to Apply: August 25, 2026" },
                { "key": "date_3", "label": "Document Verification Counseling: August 28-30, 2026" }
            ],
            "fees": [
                { "key": "fee_1", "label": "Application Processing Fee: ₹1,500" },
                { "key": "fee_2", "label": "Caution Deposit: ₹5,000 (Refundable)" },
                { "key": "fee_3", "label": f"First Year Course Fee: Approx ₹95,000 (standard program)" }
            ],
            "hostel": [
                { "key": "hostel_1", "label": "Request hostel room allocation during physical counseling" },
                { "key": "hostel_2", "label": "Submit Hostel Request form on internal dashboard" },
                { "key": "hostel_3", "label": "Pay Mess Advance & Hostel Rent deposit (₹35,000)" }
            ],
            "scholarship": [
                { "key": "scholar_1", "label": "Review available scholarships on official dashboard" },
                { "key": "scholar_2", "label": "Verify academic CGPA cut-off requirements" },
                { "key": "scholar_3", "label": "Submit Scholarship Claim form with parent income proof" }
            ]
        }
        
    return jsonify({
        'university_id': univ_id,
        'university_name': univ.university_name,
        'checklist': checklist_data,
        'completed_steps': completed_steps
    })

@app.route('/api/student/admission/checklist/toggle', methods=['POST'])
@token_required
def admission_checklist_toggle():
    """Toggles checklist step completion status and updates database."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    univ_id = body.get('university_id', '').strip()
    step_key = body.get('step_key', '').strip()
    
    if not univ_id or not step_key:
        return jsonify({'success': False, 'message': 'University ID and Step Key are required.'}), 400
        
    row = AdmissionChecklist.query.filter_by(user_id=user_email, university_id=univ_id).first()
    if not row:
        row = AdmissionChecklist(
            user_id=user_email,
            university_id=univ_id,
            completed_steps=json.dumps([step_key]),
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        db.session.add(row)
    else:
        try:
            steps = json.loads(row.completed_steps)
        except Exception:
            steps = []
            
        if step_key in steps:
            steps.remove(step_key)
        else:
            steps.append(step_key)
            
        row.completed_steps = json.dumps(steps)
        row.timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
    db.session.commit()
    return jsonify({
        'success': True,
        'completed_steps': json.loads(row.completed_steps)
    })

@app.route('/api/student/interview/start', methods=['POST'])
@token_required
def interview_start():
    """Starts a new mock interview session."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    track = body.get('track', 'AI').strip()
    
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    first_question = "Explain the difference between supervised and unsupervised learning."
    if api_key:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = f"Generate the first short, technical interview question for a mock interview focused on the topic: '{track}'"
            response = model.generate_content(prompt)
            first_question = response.text.strip()
        except Exception as ex:
            print(f"Gemini start question error: {ex}")
            
    MockInterviewSession.query.filter_by(user_id=user_email).delete()
    
    sess = MockInterviewSession(
        user_id=user_email,
        track=track,
        history=json.dumps([{"q": first_question, "a": ""}]),
        score=0,
        feedback='',
        timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    )
    db.session.add(sess)
    db.session.commit()
    
    return jsonify(sess.to_dict())

@app.route('/api/student/interview/submit', methods=['POST'])
@token_required
def interview_submit():
    """Grades answer, updates score and feedback structure, and outputs next question or scorecard."""
    user_email = request.user['user_id']
    body = request.get_json() or {}
    answer = body.get('answer', '').strip()
    
    sess = MockInterviewSession.query.filter_by(user_id=user_email).first()
    if not sess:
        return jsonify({'success': False, 'message': 'Interview session not found.'}), 404
        
    history = json.loads(sess.history)
    
    if history:
        history[-1]['a'] = answer
        
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    if len(history) >= 3:
        # Candidate completed all 3 questions, evaluate!
        score = 70
        feedback_obj = {
            "score": score,
            "feedback": "Interview completed successfully. Good attempt!",
            "strengths": ["Completed all mock interview questions.", "Exhibited interest in the domain."],
            "improvements": ["Elaborate further on underlying technical mechanisms.", "Give concrete project examples in answers."],
            "evaluations": []
        }
        
        for idx, item in enumerate(history):
            feedback_obj["evaluations"].append({
                "question": item["q"],
                "answer": item["a"],
                "rating": "Average",
                "critique": "Answer provides basic definition but could include practical trade-offs."
            })
            
        if api_key:
            try:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = (
                    f"Evaluate this mock interview for candidate in track '{sess.track}'.\n"
                    f"History:\n{json.dumps(history, indent=2)}\n\n"
                    f"Output a valid JSON object strictly matching this format:\n"
                    f"{{\n"
                    f"  \"score\": 85,\n"
                    f"  \"feedback\": \"Write a friendly 3-sentence summary of candidate performance.\",\n"
                    f"  \"strengths\": [\"Strength A\", \"Strength B\"],\n"
                    f"  \"improvements\": [\"Improvement A\", \"Improvement B\"],\n"
                    f"  \"evaluations\": [\n"
                    f"    {{\n"
                    f"      \"question\": \"Question text\",\n"
                    f"      \"answer\": \"Answer text\",\n"
                    f"      \"rating\": \"Excellent/Good/Average/Needs Improvement\",\n"
                    f"      \"critique\": \"Write a 1-sentence analysis of their answer.\"\n"
                    f"    }}\n"
                    f"  ]\n"
                    f"}}\n"
                    f"Do not output markdown code blocks or extra text, return only raw JSON."
                )
                response = model.generate_content(prompt)
                text = response.text.strip()
                if text.startswith("```json"):
                    text = text.replace("```json", "", 1).rstrip("`").strip()
                elif text.startswith("```"):
                    text = text.replace("```", "", 1).rstrip("`").strip()
                obj = json.loads(text)
                
                score = obj.get('score', 75)
                feedback_obj = obj
            except Exception as ex:
                print(f"Gemini evaluate interview error: {ex}")
                
        # Save structural JSON string back to DB
        sess.score = score
        sess.feedback = json.dumps(feedback_obj)
        sess.history = json.dumps(history)
        db.session.commit()
        
        return jsonify({
            'finished': True,
            'score': score,
            'feedback': feedback_obj,
            'history': history
        })
    else:
        # Continue mock interview questions
        next_question = "Explain how you handle failure or error exceptions in production environments."
        if api_key:
            try:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = (
                    f"Based on the mock interview history so far:\n"
                    f"{json.dumps(history, indent=2)}\n\n"
                    f"Generate the next short technical interview question for track '{sess.track}' that continues the conversation logically."
                )
                response = model.generate_content(prompt)
                next_question = response.text.strip()
            except Exception as ex:
                print(f"Gemini next question error: {ex}")
                
        history.append({"q": next_question, "a": ""})
        sess.history = json.dumps(history)
        db.session.commit()
        
        return jsonify({
            'finished': False,
            'next_question': next_question,
            'history': history
        })

@app.route('/api/community/threads', methods=['GET', 'POST'])
@token_required
def community_threads():
    """Fetches or submits forum threads."""
    user_email = request.user['user_id']
    user_name = request.user['fullname']
    
    if request.method == 'GET':
        category = request.args.get('category', '').strip()
        if category:
            threads = CommunityThread.query.filter_by(category=category).order_by(CommunityThread.id.desc()).all()
        else:
            threads = CommunityThread.query.order_by(CommunityThread.id.desc()).all()
        return jsonify([t.to_dict() for t in threads])
        
    elif request.method == 'POST':
        body = request.get_json() or {}
        title = body.get('title', '').strip()
        content = body.get('content', '').strip()
        category = body.get('category', 'General').strip()
        
        if not title or not content:
            return jsonify({'success': False, 'message': 'Title and content are required.'}), 400
            
        t = CommunityThread(
            user_id=user_email,
            user_name=user_name,
            title=title,
            content=content,
            category=category,
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        db.session.add(t)
        db.session.commit()
        return jsonify(t.to_dict()), 201

@app.route('/api/community/threads/<int:thread_id>/posts', methods=['GET', 'POST'])
@token_required
def community_posts(thread_id):
    """Fetches replies or posts a reply to a thread."""
    user_email = request.user['user_id']
    user_name = request.user['fullname']
    
    thread = CommunityThread.query.get(thread_id)
    if not thread:
        return jsonify({'success': False, 'message': 'Thread not found.'}), 404
        
    if request.method == 'GET':
        posts = CommunityPost.query.filter_by(thread_id=thread_id).order_by(CommunityPost.id.asc()).all()
        return jsonify([p.to_dict() for p in posts])
        
    elif request.method == 'POST':
        body = request.get_json() or {}
        content = body.get('content', '').strip()
        
        if not content:
            return jsonify({'success': False, 'message': 'Post content is required.'}), 400
            
        p = CommunityPost(
            thread_id=thread_id,
            user_id=user_email,
            user_name=user_name,
            content=content,
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        db.session.add(p)
        db.session.commit()
        return jsonify(p.to_dict()), 201

@app.route('/api/student/study-assistant', methods=['POST'])
@token_required
def study_assistant():
    """Generates study summaries, MCQs and interactive quizzes from syllabus text."""
    body = request.get_json() or {}
    syllabus_text = body.get('text', '').strip()
    
    if not syllabus_text:
        return jsonify({'success': False, 'message': 'Syllabus content text is required.'}), 400
        
    api_setting = SystemSetting.query.filter_by(key='gemini_api_key').first()
    api_key = api_setting.value if api_setting else os.environ.get("GEMINI_API_KEY")
    
    if not api_key:
        return jsonify({'success': False, 'message': 'Gemini API key is not configured.'}), 400
        
    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = (
            f"Analyze this syllabus/text notes:\n"
            f"'{syllabus_text}'\n\n"
            f"Generate a JSON structure containing:\n"
            f"1. A short summary of key concepts.\n"
            f"2. Three flashcards (question and answer keys).\n"
            f"3. Three multiple choice quiz questions (question, options, and answer keys).\n"
            f"Format strictly as JSON without extra text:\n"
            f"{{\n"
            f"  \"summary\": \"Text summary...\",\n"
            f"  \"flashcards\": [\n"
            f"    {{\"q\": \"Question A\", \"a\": \"Answer A\"}}\n"
            f"  ],\n"
            f"  \"quiz\": [\n"
            f"    {{\"q\": \"Quiz Q1\", \"options\": [\"Opt1\", \"Opt2\"], \"a\": \"Opt1\"}}\n"
            f"  ]\n"
            f"}}\n"
            f"Do not return markdown headers or formatting, return raw JSON string."
        )
        response = model.generate_content(prompt)
        text = response.text.strip()
        if text.startswith("```json"):
            text = text.replace("```json", "", 1).rstrip("`").strip()
        elif text.startswith("```"):
            text = text.replace("```", "", 1).rstrip("`").strip()
            
        parsed = json.loads(text)
        return jsonify(parsed)
    except Exception as e:
        app.logger.error(f"Study assistant generation failed: {e}")
        return jsonify({'success': False, 'message': f"Study assistant generation failed: {str(e)}"}), 500

def start_scheduler(app_instance):
    if app_instance.debug and os.environ.get('WERKZEUG_RUN_MAIN') != 'true':
        print("SCHEDULER: Skipped background thread in parent process.")
        return
    import threading
    def scheduler_loop():
        time.sleep(5)
        with app_instance.app_context():
            if not DocumentChunk.query.first():
                try:
                    app_instance.logger.info("SCHEDULER: Seeding RAG knowledge base chunks in background...")
                    rebuild_rag_document_chunks()
                except Exception as e:
                    app_instance.logger.warning(f"RAG rebuild failed: {e}")
                finally:
                    db.session.remove()
        while True:
            try:
                app_instance.logger.info("SCHEDULER: Starting periodic crawling loop...")
                from services.realtime_fetcher import scrape_university_details
                from datetime import datetime
                
                with app_instance.app_context():
                    unis = [u.id for u in University.query.all()]
                    db.session.remove()
                    
                for u_id in unis:
                    with app_instance.app_context():
                        u = University.query.get(u_id)
                        if u:
                            if u.last_updated:
                                try:
                                    dt = datetime.strptime(u.last_updated, "%Y-%m-%d %H:%M:%S")
                                    if (datetime.now() - dt).total_seconds() < 24 * 3600:
                                        db.session.remove()
                                        continue
                                except Exception:
                                    pass
                            try:
                                scrape_university_details(u.id)
                            except Exception as e:
                                db.session.rollback()
                                app_instance.logger.error(f"Scheduler sync failed for {u.id}: {e}")
                        db.session.remove()
                    time.sleep(5.0)
            except Exception as e:
                app_instance.logger.error(f"Scheduler global loop error: {e}")
            time.sleep(24 * 3600)
            
    t = threading.Thread(target=scheduler_loop, daemon=True)
    t.start()

if __name__ == '__main__':
    start_scheduler(app)
    app.run(debug=True, host='127.0.0.1', port=5000)
